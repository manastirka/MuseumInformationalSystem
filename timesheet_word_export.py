#!/usr/bin/env python3
"""
Timesheet Word Export Module for PostgreSQL
Generates Word documents matching the official museum "Radna Lista" format.
"""

import os
import json
import calendar
from datetime import datetime, date
from serbian_holidays import SerbianHolidays

try:
    from docx import Document
    from docx.shared import Pt, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENTATION, WD_SECTION_START
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

import psycopg
from psycopg.rows import dict_row


def _set_cell_shading(cell, color_hex):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def _set_cell_borders(cell, top=True, bottom=True, left=True, right=True, sz='4', color='000000'):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, enabled in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if enabled:
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), sz)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), color)
            tcBorders.append(border)
    tcPr.append(tcBorders)


def _set_cell_vertical_alignment(cell, val='center'):
    """Set cell vertical alignment."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), val)
    tcPr.append(vAlign)


def _set_row_height(row, height_twips, rule='exact'):
    """Set row height."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height_twips))
    trHeight.set(qn('w:hRule'), rule)
    trPr.append(trHeight)


def _merge_cells_in_row(table, row_idx, start_col, end_col):
    """Merge cells in a row from start_col to end_col (inclusive)."""
    cell_start = table.cell(row_idx, start_col)
    cell_end = table.cell(row_idx, end_col)
    cell_start.merge(cell_end)
    return cell_start


def _add_cell_text(cell, text, font_size=9, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Add text to a cell with formatting."""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(str(text))
    run.font.size = Pt(font_size)
    run.font.bold = bold
    # Set paragraph spacing to 0
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def _set_cell_width(cell, width_cm):
    """Set cell width with exact sizing."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove any existing tcW
    for existing in tcPr.findall(qn('w:tcW')):
        tcPr.remove(existing)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def _disable_table_autofit(table):
    """Disable table autofit so column widths are respected."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    # Set fixed layout
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    # Set minimal cell margins for the whole table
    tblCellMar = OxmlElement('w:tblCellMar')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), '20')  # ~0.035cm minimal padding
        el.set(qn('w:type'), 'dxa')
        tblCellMar.append(el)
    tblPr.append(tblCellMar)


def _vmerge_cell(cell, restart=False):
    """Set vertical merge on a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vMerge = OxmlElement('w:vMerge')
    if restart:
        vMerge.set(qn('w:val'), 'restart')
    tcPr.append(vMerge)


def generate_word_document(report_id, database_url):
    """
    Generate Word document matching the official museum Radna Lista format.

    Args:
        report_id: ID of the timesheet report
        database_url: PostgreSQL connection string

    Returns:
        Path to generated Word document
    """
    if not DOCX_AVAILABLE:
        raise Exception("Word export није доступан. Инсталирајте python-docx: pip install python-docx")

    pg_url = database_url.replace('postgresql+psycopg://', 'postgresql://')
    serbian_holidays = SerbianHolidays()

    # Fetch report data
    with psycopg.connect(conninfo=pg_url, row_factory=dict_row) as conn:
        with conn.cursor() as cur:
            cur.execute("""
                SELECT id, employee_name, month, year,
                       organization_unit, position,
                       extraordinary_tasks, duties_summary,
                       created_at
                FROM timesheet_reports
                WHERE id = %s
            """, (report_id,))
            header = cur.fetchone()
            if not header:
                raise Exception(f"Извештај {report_id} није пронађен")

            # Get daily data
            cur.execute("""
                SELECT day, work_in_museum, work_outside,
                       vacation, public_holiday, paid_leave,
                       other_leave, sick_leave_lt30, sick_leave_gte30
                FROM timesheet_report_days
                WHERE report_id = %s
                ORDER BY day
            """, (report_id,))
            daily_rows = cur.fetchall()

    # Build daily data dict
    daily_data = {}
    for row in daily_rows:
        daily_data[row['day']] = row

    month = header['month']
    year = header['year']
    days_in_month = calendar.monthrange(year, month)[1]

    month_names = [
        "", "јануар", "фебруар", "март", "април", "мај", "јун",
        "јул", "август", "септембар", "октобар", "новембар", "децембар"
    ]

    # =====================================================
    # PAGE 1: LANDSCAPE - Main timesheet table
    # =====================================================
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(0.5)

    # The entire page 1 is ONE big table
    # Columns: Label(col0) + 31 day columns(col1-31) + Укупно(col32)
    # = 33 columns total
    num_cols = 33

    # Total rows:
    # 0: Museum name (merged)
    # 1: Header info row 1 (Име и презиме / РАДНА ЛИСТА / Организациона јединица)
    # 2: Header info row 2 (Name value / Радно место / Department)
    # 3: Датум row (day numbers)
    # 4: Радни сати у музеју
    # 5: Ван музеја
    # 6: Годишњи одмор
    # 7: Државни празник
    # 8: Плаћено одсуство
    # 9: Остало одсуство
    # 10: Боловање < 30 дана
    # 11: Боловање > 30 дана
    # 12: Grand total row
    # 13: separator
    # 14: Повећан обим послова (label + content)
    # 15-17: continuation rows
    # 18: separator
    # 19: Ванредни послови + employee name
    # 20: continuation
    # 21: Потпис запосленог
    # 22: continuation
    # 23: separator
    # 24: Напомена шефа
    # 25-26: continuation
    # 27: Потпис шефа
    # 28-29: continuation
    # 30: separator
    # 31: Потпис директора + Повећање/умањење
    # 32: continuation with +/-
    num_rows = 33

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _disable_table_autofit(table)

    # Available width: 29.7 - 1.7 - 0.5 = 27.5cm
    label_col_width = 4.2  # cm for first column (wide enough for "Боловање < 30 дана")
    total_col_width = 1.8  # cm for total column
    day_col_width = (27.5 - label_col_width - total_col_width) / 31  # ~0.69cm per day

    # Set grid column widths in tblGrid
    tblGrid = table._tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        for i, gridCol in enumerate(tblGrid.findall(qn('w:gridCol'))):
            if i == 0:
                gridCol.set(qn('w:w'), str(int(label_col_width * 567)))
            elif i == num_cols - 1:
                gridCol.set(qn('w:w'), str(int(total_col_width * 567)))
            else:
                gridCol.set(qn('w:w'), str(int(day_col_width * 567)))

    # --- ROW 0: Museum name (merge all columns) ---
    merged = _merge_cells_in_row(table, 0, 0, num_cols - 1)
    _add_cell_text(merged, 'Природњачки музеј у Београду', font_size=16, bold=True)
    _set_row_height(table.rows[0], 712)

    # --- ROW 1: Header info ---
    # Merge cols 0-8 for "Име и презиме: [placeholder]"
    cell_name_label = _merge_cells_in_row(table, 1, 0, 8)
    p = cell_name_label.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Име и презиме: ')
    run.font.size = Pt(9)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # Merge cols 9-25 for "РАДНА ЛИСТА за месец [month] [year]."
    cell_title = _merge_cells_in_row(table, 1, 9, 25)
    p = cell_title.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'РАДНА ЛИСТА за месец  {month_names[month]} {year}.')
    run.font.size = Pt(9)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # Merge cols 26-32 for "Организациона јединица:"
    cell_org_label = _merge_cells_in_row(table, 1, 26, num_cols - 1)
    p = cell_org_label.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Организациона јединица:')
    run.font.size = Pt(9)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    _set_row_height(table.rows[1], 266)

    # --- ROW 2: Values ---
    cell_name_val = _merge_cells_in_row(table, 2, 0, 8)
    _add_cell_text(cell_name_val, header['employee_name'], font_size=13, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    cell_position = _merge_cells_in_row(table, 2, 9, 25)
    p = cell_position.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Радно место: ')
    run.font.size = Pt(9)
    run2 = p.add_run(header.get('position') or '')
    run2.font.size = Pt(10)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    cell_org_val = _merge_cells_in_row(table, 2, 26, num_cols - 1)
    _add_cell_text(cell_org_val, header.get('organization_unit') or '', font_size=13, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    _set_row_height(table.rows[2], 393)

    # --- ROW 3: Day numbers header ---
    _add_cell_text(table.cell(3, 0), 'Датум', font_size=9, bold=False)
    for day in range(1, 32):
        col = day  # col 1 = day 1, col 31 = day 31
        if day <= days_in_month:
            _add_cell_text(table.cell(3, col), str(day), font_size=9)
        else:
            _add_cell_text(table.cell(3, col), '', font_size=9)
    _add_cell_text(table.cell(3, 32), 'Укупно:', font_size=9)
    _set_row_height(table.rows[3], 281)

    # --- ROWS 4-11: Data categories ---
    categories = [
        ('Радни сати у музеју', 'work_in_museum'),
        ('Ван музеја', 'work_outside'),
        ('Годишњи одмор', 'vacation'),
        ('Државни празник', 'public_holiday'),
        ('Плаћено одсуство', 'paid_leave'),
        ('Остало одсуство', 'other_leave'),
        ('Боловање < 30 дана', 'sick_leave_lt30'),
        ('Боловање > 30 дана', 'sick_leave_gte30'),
    ]

    for cat_idx, (label, field) in enumerate(categories):
        row_idx = 4 + cat_idx
        _add_cell_text(table.cell(row_idx, 0), label, font_size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT)

        row_total = 0
        for day in range(1, 32):
            col = day
            cell = table.cell(row_idx, col)

            if day > days_in_month:
                _add_cell_text(cell, '', font_size=9)
                continue

            day_date = date(year, month, day)
            weekday = day_date.weekday()
            is_weekend = weekday >= 5
            holiday = serbian_holidays.is_holiday(day_date)

            day_data = daily_data.get(day, {})
            value = float(day_data.get(field, 0) or 0)

            # For "Радни сати у музеју" row, show "Х" for weekends/holidays with no data
            if field == 'work_in_museum' and (is_weekend or holiday) and value == 0:
                _add_cell_text(cell, 'Х', font_size=9)
            elif value > 0:
                _add_cell_text(cell, str(int(value)), font_size=9)
                row_total += value
            else:
                _add_cell_text(cell, '', font_size=9)

        # Total column
        total_text = str(int(row_total)) if row_total > 0 else ''
        _add_cell_text(table.cell(row_idx, 32), total_text, font_size=9)
        _set_row_height(table.rows[row_idx], 281)

    # --- ROW 12: Grand total row ---
    # Merge cols 0-31 (empty), col 32 = grand total
    merged_total = _merge_cells_in_row(table, 12, 0, 31)
    _add_cell_text(merged_total, '', font_size=9)

    # Calculate grand total
    grand_total = 0
    for day in range(1, days_in_month + 1):
        day_data = daily_data.get(day, {})
        for _, field in categories:
            grand_total += float(day_data.get(field, 0) or 0)

    _add_cell_text(table.cell(12, 32), str(int(grand_total)) if grand_total > 0 else '', font_size=9)
    _set_row_height(table.rows[12], 281)

    # --- ROW 13: Separator ---
    merged_sep1 = _merge_cells_in_row(table, 13, 0, num_cols - 1)
    _add_cell_text(merged_sep1, '', font_size=2)
    _set_row_height(table.rows[13], 133)

    # --- ROW 14: Повећан обим послова ---
    cell_label14 = _merge_cells_in_row(table, 14, 0, 3)
    _add_cell_text(cell_label14, 'Повећан обим послова', font_size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    cell_val14 = _merge_cells_in_row(table, 14, 4, num_cols - 1)
    _add_cell_text(cell_val14, '', font_size=9)
    _set_row_height(table.rows[14], 394)

    # Rows 15-17: continuation (merged single cells)
    for r in [15, 16, 17]:
        merged = _merge_cells_in_row(table, r, 0, num_cols - 1)
        _add_cell_text(merged, '', font_size=9)
        height = 394 if r == 15 else 422
        _set_row_height(table.rows[r], height)

    # --- ROW 18: Separator ---
    merged_sep2 = _merge_cells_in_row(table, 18, 0, num_cols - 1)
    _add_cell_text(merged_sep2, '', font_size=2)
    _set_row_height(table.rows[18], 133)

    # --- ROW 19: Ванредни послови + employee name ---
    cell_label19 = _merge_cells_in_row(table, 19, 0, 3)
    _add_cell_text(cell_label19, 'Ванредни послови', font_size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    cell_val19 = _merge_cells_in_row(table, 19, 4, num_cols - 1)
    _add_cell_text(cell_val19, '', font_size=9)
    _set_row_height(table.rows[19], 372)

    # Row 20: continuation
    cell_label20 = _merge_cells_in_row(table, 20, 0, 3)
    _add_cell_text(cell_label20, '', font_size=9)
    cell_val20 = _merge_cells_in_row(table, 20, 4, num_cols - 1)
    # Employee name (last, first)
    name_parts = (header['employee_name'] or '').split(' ')
    reversed_name = ' '.join(reversed(name_parts)) if len(name_parts) > 1 else header['employee_name']
    _add_cell_text(cell_val20, reversed_name, font_size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    _set_row_height(table.rows[20], 422)

    # Row 21: Потпис запосленог
    cell_label21 = _merge_cells_in_row(table, 21, 0, 3)
    _add_cell_text(cell_label21, 'Потпис запосленог', font_size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    cell_val21 = _merge_cells_in_row(table, 21, 4, num_cols - 1)
    _add_cell_text(cell_val21, '', font_size=9)
    _set_row_height(table.rows[21], 398)

    # Row 22: continuation
    cell_label22 = _merge_cells_in_row(table, 22, 0, 3)
    _add_cell_text(cell_label22, '', font_size=9)
    cell_val22 = _merge_cells_in_row(table, 22, 4, num_cols - 1)
    _add_cell_text(cell_val22, '', font_size=9)
    _set_row_height(table.rows[22], 422)

    # --- ROW 23: Separator ---
    merged_sep3 = _merge_cells_in_row(table, 23, 0, num_cols - 1)
    _add_cell_text(merged_sep3, '', font_size=2)
    _set_row_height(table.rows[23], 133)

    # --- ROW 24: Напомена шефа ---
    cell_label24 = _merge_cells_in_row(table, 24, 0, 3)
    _add_cell_text(cell_label24, 'Напомена шефа', font_size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    cell_val24 = _merge_cells_in_row(table, 24, 4, num_cols - 1)
    _add_cell_text(cell_val24, '', font_size=9)
    _set_row_height(table.rows[24], 213)

    # Rows 25-26: continuation
    for r in [25, 26]:
        cell_l = _merge_cells_in_row(table, r, 0, 3)
        _add_cell_text(cell_l, '', font_size=9)
        cell_v = _merge_cells_in_row(table, r, 4, num_cols - 1)
        _add_cell_text(cell_v, '', font_size=9)
        _set_row_height(table.rows[r], 299)

    # Row 27: Потпис шефа
    cell_label27 = _merge_cells_in_row(table, 27, 0, 3)
    _add_cell_text(cell_label27, 'Потпис шефа', font_size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    cell_val27 = _merge_cells_in_row(table, 27, 4, num_cols - 1)
    _add_cell_text(cell_val27, '', font_size=9)
    _set_row_height(table.rows[27], 230)

    # Rows 28-29: continuation
    for r in [28, 29]:
        cell_l = _merge_cells_in_row(table, r, 0, 3)
        _add_cell_text(cell_l, '', font_size=9)
        cell_v = _merge_cells_in_row(table, r, 4, num_cols - 1)
        _add_cell_text(cell_v, '', font_size=9)
        _set_row_height(table.rows[r], 264)

    # --- ROW 30: Separator ---
    merged_sep4 = _merge_cells_in_row(table, 30, 0, num_cols - 1)
    _add_cell_text(merged_sep4, '', font_size=2)
    _set_row_height(table.rows[30], 133)

    # --- ROW 31: Потпис директора + Повећање/умањење ---
    cell_director = _merge_cells_in_row(table, 31, 0, 25)
    _add_cell_text(cell_director, 'Потпис директора', font_size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    cell_salary = _merge_cells_in_row(table, 31, 26, num_cols - 1)
    _add_cell_text(cell_salary, 'Повећање / умањење зараде за обрачун', font_size=9)
    _set_row_height(table.rows[31], 299)

    # --- ROW 32: continuation with +/- ---
    cell_dir2 = _merge_cells_in_row(table, 32, 0, 25)
    _add_cell_text(cell_dir2, '', font_size=9)
    cell_plus = _merge_cells_in_row(table, 32, 26, 30)
    _add_cell_text(cell_plus, '+', font_size=9)
    cell_minus = _merge_cells_in_row(table, 32, 31, num_cols - 1)
    _add_cell_text(cell_minus, '-', font_size=9)
    _set_row_height(table.rows[32], 299)

    # Set column widths
    for row in table.rows:
        for col_idx, cell in enumerate(row.cells):
            if col_idx == 0:
                _set_cell_width(cell, label_col_width)
            elif col_idx == 32:
                _set_cell_width(cell, total_col_width)
            else:
                _set_cell_width(cell, day_col_width)

    # =====================================================
    # PAGE 2: PORTRAIT - Work description
    # =====================================================
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    new_section.orientation = WD_ORIENTATION.PORTRAIT
    new_section.page_width = Cm(21.0)
    new_section.page_height = Cm(29.7)
    new_section.top_margin = Cm(1.0)
    new_section.bottom_margin = Cm(1.0)
    new_section.left_margin = Cm(1.0)
    new_section.right_margin = Cm(1.0)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Обављени послови ')
    run.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{month_names[month]} {year}.')
    run.font.size = Pt(18)

    # Empty lines
    for _ in range(4):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('')
        run.font.size = Pt(18)

    # Work categories from extraordinary_tasks (JSON)
    work_category_labels = {
        'cat_2_1': '2.1. РАД НА ЗБИРКАМА',
        'cat_2_2': '2.2. ИДЕНТИФИКАЦИЈЕ И ЕКСПЕРТИЗЕ',
        'cat_2_3': '2.3. НАУЧНИ И СТРУЧНИ РАД',
        'cat_3_1': '3.1. ТЕРЕНСКИ РАД',
        'cat_3_2': '3.2. НОВИ ПРЕДМЕТИ У ЗБИРЦИ',
        'cat_4_1': '4.1. ТЕКУЋЕ ИЗЛОЖБЕ',
        'cat_4_2': '4.2. ПЛАНИРАНЕ ИЗЛОЖБЕ',
        'cat_5_1': '5.1. САРАДЊА СА ПЕДАГОШКОМ СЛУЖБОМ',
        'cat_5_2': '5.2. СТРУЧНО ВОЂЕЊЕ / ПРЕДАВАЊА / РАДИОНИЦЕ',
        'cat_5_3': '5.3. МЕДИЈИ / PR / КОМУНИКАЦИЈА',
        'cat_6_1': '6.1. ПРОЈЕКТИ / РАДНЕ ГРУПЕ',
        'cat_6_2': '6.2. НАЦИОНАЛНА И МЕЂУНАРОДНА ПАРТНЕРСТВА',
        'cat_7': '7. ДИГИТАЛИЗАЦИЈА И ДОКУМЕНТАЦИЈА',
        'cat_8': '8. ОСТАЛИ ПОСЛОВИ И АКТИВНОСТИ'
    }

    extraordinary_tasks = header.get('extraordinary_tasks')
    has_work_data = False

    if extraordinary_tasks:
        try:
            work_data = json.loads(extraordinary_tasks)

            # Group items by category and output
            for key, label in work_category_labels.items():
                content = work_data.get(key, '').strip()
                if content:
                    has_work_data = True
                    # Category header (bold)
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(label)
                    run.font.size = Pt(8)
                    run.font.bold = True

                    # Content lines
                    for line in content.split('\n'):
                        line = line.strip()
                        if line:
                            p = doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            run = p.add_run(line)
                            run.font.size = Pt(10)

        except (json.JSONDecodeError, TypeError):
            # Old format - plain text
            if extraordinary_tasks.strip():
                has_work_data = True
                for line in extraordinary_tasks.strip().split('\n'):
                    line = line.strip()
                    if line:
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run = p.add_run(line)
                        run.font.size = Pt(10)

    if not has_work_data:
        p = doc.add_paragraph()
        run = p.add_run('Нема података о обављеним пословима.')
        run.font.size = Pt(10)

    # Duties summary
    if header.get('duties_summary'):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header['duties_summary'])
        run.font.size = Pt(10)

    # =====================================================
    # Save document
    # =====================================================
    safe_name = header['employee_name'].replace(' ', '_').replace('.', '').replace(',', '')
    month_name_file = month_names[month] if 1 <= month <= 12 else str(month)
    filename = f"RadnaLista_{safe_name}_{month_name_file}_{year}.docx"

    output_dir = os.path.join(os.path.dirname(__file__), 'exports', 'timesheets')
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, filename)

    doc.save(output_path)
    return output_path


if __name__ == '__main__':
    import sys

    if len(sys.argv) < 3:
        print("Usage: python timesheet_word_export.py <report_id> <database_url>")
        sys.exit(1)

    report_id = int(sys.argv[1])
    database_url = sys.argv[2]

    try:
        output_path = generate_word_document(report_id, database_url)
        print(f"Документ генерисан: {output_path}")
    except Exception as e:
        print(f"Грешка: {str(e)}")
        sys.exit(1)

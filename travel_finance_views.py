"""Shared route implementations for travel and finance workflows."""

import json
import logging
import math
import os
import random
from datetime import datetime, timedelta
from io import BytesIO

import requests as http_requests
from flask import jsonify, render_template, request, send_file, session

logger = logging.getLogger(__name__)

# Official K1 toll tariffs from the Putevi Srbije PDF price list,
# effective 2026-01-01 07:00. The route presets below use a small alias layer
# for destinations that map to the nearest published toll station.
TOLL_PRICES_FROM_BELGRADE = {
    # Beograd - Presevo
    'MALI POŽAREVAC': 100,
    'UMČARI': 140,
    'VODANJ': 170,
    'KOLARI': 200,
    'SMEDEREVO': 230,
    'POŽAREVAC': 240,
    'VELIKA PLANA': 400,
    'MARKOVAC': 460,
    'LAPOVO': 500,
    'BATOČINA': 520,
    'JAGODINA': 640,
    'ĆUPRIJA': 710,
    'PARAĆIN': 770,
    'ĆIĆEVAC': 860,
    'KRUŠEVAC ISTOK': 960,
    'KRUŠEVAC ZAPAD': 970,
    'KOŠEVI': 1000,
    'VELIKA DRENOVA': 1050,
    'TRSTENIK': 1120,
    'VRNJAČKA BANJA': 1190,
    'VRBA': 1260,
    'RAŽANJ': 900,
    'ALEKSINAČKI RUDNICI': 1020,
    'ALEKSINAC': 1050,
    'NIŠ JUG': 1180,
    'MEROŠINA': 1210,
    'DOLJEVAC SELO': 1270,
    'BRESTOVAC': 1300,
    'LESKOVAC CENTAR': 1340,
    'LESKOVAC JUG': 1430,
    'GRDELICA': 1490,
    'PREDEJANE': 1580,
    'VLADIČIN HAN': 1720,
    'VRANJE': 1850,
    'BUJANOVAC SEVER': 1920,
    'BUJANOVAC JUG': 1940,
    'PREŠEVO': 2060,
    # Beograd - Dimitrovgrad, derived from the official Dimitrovgrad-Beograd matrix
    'BELA PALANKA': 1430,
    'PIROT ZAPAD': 1570,
    'PIROT ISTOK': 1640,
    'DIMITROVGRAD': 1800,
    # Beograd - Šid
    'ADAŠEVCI': 50,
    'MOROVIĆ': 50,
    'KUZMIN': 120,
    'SREMSKA MITROVICA': 240,
    'RUMA': 320,
    'HRTKOVCI': 380,
    'ŠABAC': 460,
    'PEĆINCI': 380,
    'ŠIMANOVCI': 520,
    # Beograd - Subotica
    'STARA PAZOVA': 100,
    'INĐIJA': 100,
    'MARADIK': 120,
    'BEŠKA': 150,
    'KOVILJ': 230,
    'NOVI SAD': 340,
    'ZMAJEVO': 430,
    'VRBAS': 500,
    'FEKETIĆ': 550,
    'BAČKA TOPOLA': 670,
    'ŽEDNIK': 730,
    'SUBOTICA': 850,
    # Beograd - Požega
    'UB': 150,
    'LAJKOVAC': 220,
    'LJIG': 360,
    'TAKOVO': 540,
    'PRELJINA': 640,
    'ČAČAK (PAKOVRAĆE)': 720,
    'LUČANI': 860,
    'PRILIPAC': 950,
    # Aliases used by the destination presets in this app
    'KRUŠEVAC': 970,
    'NIŠ': 1180,
    'DOLJEVAC': 1270,
    'LESKOVAC': 1430,
    'BUJANOVAC': 1940,
    'PIROT': 1570,
    'ČAČAK': 720,
}

DESTINATION_COORDS = {
    'STARA PLANINA': {'lat': 43.3667, 'lon': 22.6000, 'km': 320, 'toll_station': 'PIROT'},
    'ĐERDAP': {'lat': 44.6167, 'lon': 21.8833, 'km': 240, 'toll_station': 'POŽAREVAC'},
    'TARA': {'lat': 43.8833, 'lon': 19.5500, 'km': 180, 'toll_station': 'ČAČAK'},
    'ZLATIBOR': {'lat': 43.7333, 'lon': 19.7000, 'km': 230, 'toll_station': 'ČAČAK'},
    'KOPAONIK': {'lat': 43.2833, 'lon': 20.8167, 'km': 290, 'toll_station': 'ČAČAK'},
    'FRUŠKA GORA': {'lat': 45.1500, 'lon': 19.7833, 'km': 90, 'toll_station': 'NOVI SAD'},
    'DIVČIBARE': {'lat': 44.1167, 'lon': 19.9833, 'km': 120, 'toll_station': 'LAJKOVAC'},
    'GOLIJA': {'lat': 43.3667, 'lon': 20.3167, 'km': 260, 'toll_station': 'ČAČAK'},
    'NIŠ': {'lat': 43.3211, 'lon': 21.8958, 'km': 240, 'toll_station': 'NIŠ'},
    'NOVI SAD': {'lat': 45.2517, 'lon': 19.8369, 'km': 95, 'toll_station': 'NOVI SAD'},
    'SUBOTICA': {'lat': 46.1000, 'lon': 19.6667, 'km': 190, 'toll_station': 'SUBOTICA'},
    'VRANJE': {'lat': 42.5500, 'lon': 21.9000, 'km': 340, 'toll_station': 'VRANJE'},
    'PIROT': {'lat': 43.1533, 'lon': 22.5856, 'km': 310, 'toll_station': 'PIROT'},
    'LESKOVAC': {'lat': 42.9983, 'lon': 21.9461, 'km': 280, 'toll_station': 'LESKOVAC'},
    'ČAČAK': {'lat': 43.8914, 'lon': 20.3497, 'km': 150, 'toll_station': 'ČAČAK'},
    'KRUŠEVAC': {'lat': 43.5833, 'lon': 21.3333, 'km': 200, 'toll_station': 'KRUŠEVAC'},
    'VRNJAČKA BANJA': {'lat': 43.6167, 'lon': 20.9000, 'km': 210, 'toll_station': 'VRNJAČKA BANJA'},
    'JAGODINA': {'lat': 43.9833, 'lon': 21.2500, 'km': 140, 'toll_station': 'JAGODINA'},
    'PARAĆIN': {'lat': 43.8667, 'lon': 21.4000, 'km': 160, 'toll_station': 'PARAĆIN'},
    'ŠABAC': {'lat': 44.7500, 'lon': 19.6833, 'km': 90, 'toll_station': 'ŠABAC'},
    'SREMSKA MITROVICA': {'lat': 44.9667, 'lon': 19.6167, 'km': 75, 'toll_station': 'SREMSKA MITROVICA'},
    'VRBAS': {'lat': 45.5667, 'lon': 19.6333, 'km': 140, 'toll_station': 'VRBAS'},
    'ZAJEČAR': {'lat': 43.9000, 'lon': 22.2833, 'km': 250, 'toll_station': 'PARAĆIN'},
    'BOR': {'lat': 44.0833, 'lon': 22.1000, 'km': 240, 'toll_station': 'PARAĆIN'},
    'KLADOVO': {'lat': 44.6167, 'lon': 22.6167, 'km': 260, 'toll_station': 'POŽAREVAC'},
    'NEGOTIN': {'lat': 44.2333, 'lon': 22.5333, 'km': 260, 'toll_station': 'POŽAREVAC'},
    'SJENICA': {'lat': 43.2667, 'lon': 20.0000, 'km': 280, 'toll_station': 'ČAČAK'},
    'NOVI PAZAR': {'lat': 43.1333, 'lon': 20.5167, 'km': 290, 'toll_station': 'ČAČAK'},
    'UŽICE': {'lat': 43.8500, 'lon': 19.8500, 'km': 190, 'toll_station': 'ČAČAK'},
    'VALJEVO': {'lat': 44.2667, 'lon': 19.8833, 'km': 100, 'toll_station': 'LAJKOVAC'},
    'ARANĐELOVAC': {'lat': 44.3000, 'lon': 20.5500, 'km': 75, 'toll_station': None},
    'KRAGUJEVAC': {'lat': 44.0167, 'lon': 20.9167, 'km': 140, 'toll_station': 'BATOČINA'},
    'SMEDEREVO': {'lat': 44.6667, 'lon': 20.9333, 'km': 50, 'toll_station': 'SMEDEREVO'},
    'POŽAREVAC': {'lat': 44.6167, 'lon': 21.1833, 'km': 80, 'toll_station': 'POŽAREVAC'},
    'DIMITROVGRAD': {'lat': 43.0167, 'lon': 22.7833, 'km': 350, 'toll_station': 'DIMITROVGRAD'},
    'BEOGRAD': {'lat': 44.8167, 'lon': 20.4667, 'km': 0, 'toll_station': None},
}

FUEL_PRICES_LOCAL = {
    'Бензин': 174,
    'Дизел': 191,
    'BMB 95': 174,
    'BMB 100': 189,
    'Евродизел': 191,
}

# GPS coordinates of toll plazas (lat, lon) for route geometry matching.
# Positions approximate the highway interchange, not the city centre.
TOLL_STATION_LOCATIONS = {
    # Beograd - Preševo (E-75 South)
    'MALI POŽAREVAC': (44.58, 20.72),
    'UMČARI': (44.53, 20.78),
    'VODANJ': (44.50, 20.82),
    'KOLARI': (44.47, 20.87),
    'SMEDEREVO': (44.60, 20.93),
    'POŽAREVAC': (44.60, 21.18),
    'VELIKA PLANA': (44.18, 21.08),
    'MARKOVAC': (44.10, 21.07),
    'LAPOVO': (44.00, 21.10),
    'BATOČINA': (43.96, 21.08),
    'JAGODINA': (43.98, 21.26),
    'ĆUPRIJA': (43.93, 21.37),
    'PARAĆIN': (43.87, 21.40),
    'ĆIĆEVAC': (43.72, 21.45),
    'KRUŠEVAC ISTOK': (43.62, 21.38),
    'KRUŠEVAC ZAPAD': (43.56, 21.33),
    'KOŠEVI': (43.52, 21.28),
    'RAŽANJ': (43.67, 21.55),
    'ALEKSINAČKI RUDNICI': (43.58, 21.65),
    'ALEKSINAC': (43.54, 21.70),
    'NIŠ JUG': (43.28, 21.89),
    'MEROŠINA': (43.25, 21.72),
    'DOLJEVAC SELO': (43.20, 21.82),
    'BRESTOVAC': (43.16, 21.88),
    'LESKOVAC CENTAR': (43.00, 21.95),
    'LESKOVAC JUG': (42.93, 21.95),
    'GRDELICA': (42.87, 21.95),
    'PREDEJANE': (42.82, 21.97),
    'VLADIČIN HAN': (42.70, 22.00),
    'VRANJE': (42.55, 21.90),
    'BUJANOVAC SEVER': (42.47, 21.77),
    'BUJANOVAC JUG': (42.43, 21.76),
    'PREŠEVO': (42.31, 21.95),
    # Beograd - Dimitrovgrad (branches from E-75 at Niš)
    'BELA PALANKA': (43.22, 22.31),
    'PIROT ZAPAD': (43.16, 22.50),
    'PIROT ISTOK': (43.15, 22.62),
    'DIMITROVGRAD': (43.02, 22.78),
    # Beograd - Šid (E-70 West)
    'ADAŠEVCI': (45.03, 19.27),
    'MOROVIĆ': (44.95, 19.22),
    'KUZMIN': (44.97, 19.35),
    'SREMSKA MITROVICA': (44.97, 19.61),
    'RUMA': (45.01, 19.84),
    'HRTKOVCI': (44.94, 19.93),
    'ŠABAC': (44.75, 19.69),
    'PEĆINCI': (44.90, 20.00),
    'ŠIMANOVCI': (44.85, 20.10),
    # Beograd - Subotica (E-75 North)
    'STARA PAZOVA': (44.98, 20.16),
    'INĐIJA': (45.05, 20.09),
    'MARADIK': (45.10, 19.97),
    'BEŠKA': (45.12, 19.92),
    'KOVILJ': (45.20, 19.85),
    'NOVI SAD': (45.25, 19.84),
    'ZMAJEVO': (45.47, 19.68),
    'VRBAS': (45.57, 19.63),
    'FEKETIĆ': (45.67, 19.57),
    'BAČKA TOPOLA': (45.82, 19.62),
    'ŽEDNIK': (45.92, 19.68),
    'SUBOTICA': (46.07, 19.67),
    # Beograd - Požega (via Obrenovac)
    'UB': (44.47, 20.07),
    'LAJKOVAC': (44.37, 20.17),
    'LJIG': (44.22, 20.24),
    'TAKOVO': (44.05, 20.28),
    'PRELJINA': (43.92, 20.30),
    'ČAČAK (PAKOVRAĆE)': (43.89, 20.35),
    'LUČANI': (43.72, 20.28),
    'PRILIPAC': (43.62, 20.15),
}

# Aliases that map to canonical station names with known coordinates
TOLL_STATION_ALIASES = {
    'KRUŠEVAC': 'KRUŠEVAC ZAPAD',
    'NIŠ': 'NIŠ JUG',
    'DOLJEVAC': 'DOLJEVAC SELO',
    'LESKOVAC': 'LESKOVAC JUG',
    'BUJANOVAC': 'BUJANOVAC JUG',
    'PIROT': 'PIROT ZAPAD',
    'ČAČAK': 'ČAČAK (PAKOVRAĆE)',
    'VRNJAČKA BANJA': 'KOŠEVI',
    'VELIKA DRENOVA': 'KOŠEVI',
    'TRSTENIK': 'KOŠEVI',
}


def _point_near(lat1, lon1, lat2, lon2, radius_km=5):
    """Check if two points are within radius_km using fast approximation."""
    dlat = (lat1 - lat2) * 111.0
    dlon = (lon1 - lon2) * 80.0  # ~cos(44°) * 111
    return (dlat * dlat + dlon * dlon) < radius_km * radius_km


def _find_tolls_from_geometry(geometry_coords):
    """Determine which toll stations an OSM route passes through.

    Args:
        geometry_coords: list of [lon, lat] pairs from GeoJSON LineString

    Returns:
        set of toll station names that the route passes near
    """
    matched = set()
    # Sample every 3rd point for performance (route coords can be dense)
    sampled = geometry_coords[::3]
    for station_name, (slat, slon) in TOLL_STATION_LOCATIONS.items():
        for coord in sampled:
            rlon, rlat = coord[0], coord[1]
            if _point_near(rlat, rlon, slat, slon, radius_km=5):
                matched.add(station_name)
                break
    return matched


def _resolve_station_price(station_name):
    """Look up the toll price for a station, handling aliases."""
    price = TOLL_PRICES_FROM_BELGRADE.get(station_name)
    if price is not None:
        return station_name, price
    canonical = TOLL_STATION_ALIASES.get(station_name)
    if canonical:
        return canonical, TOLL_PRICES_FROM_BELGRADE.get(canonical, 0)
    return station_name, 0


def _toll_from_geometry(geometry_coords):
    """Calculate toll from OSM route geometry.

    Returns:
        (total_toll, toll_stations_set, per_station_details)
        where total_toll is the one-way toll (furthest station per highway section).
    """
    matched = _find_tolls_from_geometry(geometry_coords)
    if not matched:
        return 0, set(), []

    # Group matched stations by highway section to find the furthest on each
    # Section is identified by the price table origin
    # Stations on each section sorted by price (= distance from Belgrade)
    sections = {
        'BG_PRESEVO': [
            'MALI POŽAREVAC', 'UMČARI', 'VODANJ', 'KOLARI', 'SMEDEREVO',
            'POŽAREVAC', 'VELIKA PLANA', 'MARKOVAC', 'LAPOVO', 'BATOČINA',
            'JAGODINA', 'ĆUPRIJA', 'PARAĆIN', 'ĆIĆEVAC',
            'KRUŠEVAC ISTOK', 'KRUŠEVAC ZAPAD', 'KOŠEVI',
            'RAŽANJ', 'ALEKSINAČKI RUDNICI', 'ALEKSINAC',
            'NIŠ JUG', 'MEROŠINA', 'DOLJEVAC SELO', 'BRESTOVAC',
            'LESKOVAC CENTAR', 'LESKOVAC JUG', 'GRDELICA', 'PREDEJANE',
            'VLADIČIN HAN', 'VRANJE', 'BUJANOVAC SEVER', 'BUJANOVAC JUG',
            'PREŠEVO',
        ],
        'BG_DIMITROVGRAD': [
            'BELA PALANKA', 'PIROT ZAPAD', 'PIROT ISTOK', 'DIMITROVGRAD',
        ],
        'BG_SID': [
            'ADAŠEVCI', 'MOROVIĆ', 'KUZMIN', 'SREMSKA MITROVICA', 'RUMA',
            'HRTKOVCI', 'ŠABAC', 'PEĆINCI', 'ŠIMANOVCI',
        ],
        'BG_SUBOTICA': [
            'STARA PAZOVA', 'INĐIJA', 'MARADIK', 'BEŠKA', 'KOVILJ',
            'NOVI SAD', 'ZMAJEVO', 'VRBAS', 'FEKETIĆ',
            'BAČKA TOPOLA', 'ŽEDNIK', 'SUBOTICA',
        ],
        'BG_POZEGA': [
            'UB', 'LAJKOVAC', 'LJIG', 'TAKOVO', 'PRELJINA',
            'ČAČAK (PAKOVRAĆE)', 'LUČANI', 'PRILIPAC',
        ],
    }

    total_toll = 0
    toll_stations = set()

    for section_name, section_stations in sections.items():
        # Find stations on this section that the route passes through
        on_route = [s for s in section_stations if s in matched]
        if not on_route:
            continue
        # Furthest station = last in the ordered list
        furthest = on_route[-1]
        _, price = _resolve_station_price(furthest)
        if price > 0:
            total_toll += price
            toll_stations.add(furthest)

    return total_toll, toll_stations, []


FINANCIAL_PLAN_CATEGORY_NAMES = {
    'akvizicija_teren': '1. АКВИЗИЦИЈА - НАБАВКА МУЗЕЈСКИХ ПРЕДМЕТА НА ТЕРЕНУ',
    'akvizicija_otkup': '2. АКВИЗИЦИЈА - НАБАВКА МУЗЕЈСКИХ ПРЕДМЕТА ОТКУПОМ',
    'prepariranje': '3. ПРЕПАРИРАЊЕ И КОНЗЕРВАЦИЈА',
    'odrzavanje': '4. ОДРЖАВАЊЕ ЗБИРКИ И ДЕПОА',
    'oprema': '5. НАБАВКА НЕОПХОДНЕ ОПРЕМЕ',
    'skupovi': '6. СТРУЧНИ И НАУЧНИ СКУПОВИ',
    'projekti': '7. РЕАЛИЗАЦИЈА СТРУЧНИХ И НАУЧНИХ ПРОЈЕКАТА',
    'izlozbe': '8. ИЗЛОЖБЕНА ДЕЛАТНОСТ',
    'izdavacka': '9. ИЗДАВАЧКА ДЕЛАТНОСТ',
    'literatura': '10. НАБАВКА ЛИТЕРАТУРЕ И ПЕРИОДИКЕ',
    'usavrsavanje': '11. УСАВРШАВАЊЕ И САРАДЊА',
}

FINANCIAL_PLAN_CATEGORY_DEFINITIONS = [
    {
        'key': 'akvizicija_teren',
        'number': '1',
        'title': 'АКВИЗИЦИЈА - НАБАВКА МУЗЕЈСКИХ ПРЕДМЕТА НА ТЕРЕНУ, ТЕРЕНСКА ИСТРАЖИВАЊА',
        'activity_label': 'Активност',
    },
    {
        'key': 'akvizicija_otkup',
        'number': '2',
        'title': 'АКВИЗИЦИЈА - НАБАВКА МУЗЕЈСКИХ ПРЕДМЕТА ОТКУПОМ',
        'activity_label': 'Активност',
    },
    {
        'key': 'prepariranje',
        'number': '3',
        'title': 'ПРЕПАРИРАЊЕ И КОНЗЕРВАЦИЈА, СТРУЧНА И НАУЧНА ОБРАДА МУЗЕЈСКИХ ПРЕДМЕТА',
        'activity_label': 'Активност и потребна средства и материјал',
    },
    {
        'key': 'odrzavanje',
        'number': '4',
        'title': 'ОДРЖАВАЊЕ ЗБИРКИ И ДЕПОА, СМЕШТАЈ И ЗАШТИТА МУЗЕЈСКИХ ПРЕДМЕТА И ЗБИРКИ',
        'activity_label': 'Активност, спецификација потребних средстава',
    },
    {
        'key': 'oprema',
        'number': '5',
        'title': 'НАБАВКА НЕОПХОДНЕ ОПРЕМЕ',
        'activity_label': 'Спецификација опреме',
    },
    {
        'key': 'skupovi',
        'number': '6',
        'title': 'СТРУЧНИ И НАУЧНИ СКУПОВИ',
        'activity_label': 'Активност',
    },
    {
        'key': 'projekti',
        'number': '7',
        'title': 'РЕАЛИЗАЦИЈА СТРУЧНИХ И НАУЧНИХ ПРОЈЕКАТА',
        'activity_label': 'Активност',
        'note': 'Потребно је уз сваку планирану ставку прецизирати да ли је реч о пројектима који су раније планирани и започети и за чију реализацију су већ планирана и/или добијена средства или се ради о новим пројектима.',
    },
    {
        'key': 'izlozbe',
        'number': '8',
        'title': 'ИЗЛОЖБЕНА ДЕЛАТНОСТ',
        'activity_label': 'Спецификација трошкова',
        'has_exhibition_fields': True,
    },
    {
        'key': 'izdavacka',
        'number': '9',
        'title': 'ИЗДАВАЧКА ДЕЛАТНОСТ (ПРИПРЕМА И ИЗРАДА ПУБЛИКАЦИЈА)',
        'activity_label': 'Активност, средства',
    },
    {
        'key': 'literatura',
        'number': '10',
        'title': 'НАБАВКА НАУЧНЕ И СТРУЧНЕ ЛИТЕРАТУРЕ И ПЕРИОДИКЕ',
        'activity_label': 'Спецификација',
    },
    {
        'key': 'usavrsavanje',
        'number': '11',
        'title': 'УСАВРШАВАЊЕ, СТУДИЈСКИ БОРАВЦИ, САРАДЊА СА ИНСТИТУЦИЈАМА И/ИЛИ ПОЈЕДИНЦИМА',
        'activity_label': 'Активност',
    },
]


def _financial_plan_year_options():
    start_year = datetime.now().year + 1
    return [str(start_year + offset) for offset in range(3)]


def _normalize_financial_plan_payload(payload):
    if not isinstance(payload, dict):
        return None, {}

    if 'years' in payload and isinstance(payload.get('years'), dict):
        selected_year = str(payload.get('selectedYear') or '').strip() or None
        return selected_year, payload.get('years', {})

    return None, payload


def _financial_plan_year_total(year_data):
    if not isinstance(year_data, dict):
        return 0

    total = 0
    for category_key in FINANCIAL_PLAN_CATEGORY_NAMES:
        items = year_data.get(category_key, [])
        if not isinstance(items, list):
            continue
        for item in items:
            if isinstance(item, dict):
                total += float(item.get('amount', 0) or 0)
    return total


def render_financial_requests_page():
    return render_template('zahtevi/finansijski_zahtevi.html', page_title='Финансијски захтеви')


def render_day_off_request_page():
    return render_template(
        'zahtevi/zahtev_slobodan_dan.html',
        page_title='Захтев за слободан дан',
        user_full_name=session.get('user_name', ''),
    )


def render_vacation_request_page():
    return render_template(
        'zahtevi/zahtev_godisnji_odmor.html',
        page_title='Захтев за годишњи одмор',
        user_full_name=session.get('user_name', ''),
    )


def render_misc_request_page():
    return render_template(
        'zahtevi/zahtev_razno.html',
        page_title='Захтев разно',
        user_full_name=session.get('user_name', ''),
    )


def render_financial_plan_page():
    year_options = _financial_plan_year_options()
    selected_year = request.args.get('year', '').strip()
    if selected_year not in year_options:
        selected_year = year_options[0]

    return render_template(
        'finansije/finansijski_plan.html',
        page_title='Финансијски план',
        today=datetime.now().strftime('%Y-%m-%d'),
        year_options=year_options,
        selected_year=selected_year,
        financial_plan_categories=FINANCIAL_PLAN_CATEGORY_DEFINITIONS,
    )


def render_procurement_request_page():
    return render_template('finansije/zahtev_nabavka.html', page_title='Захтев за набавку')


def render_field_activity_page():
    activities = []
    today = datetime.now().date()

    if os.environ.get('DATABASE_URL'):
        try:
            import psycopg
            from psycopg.rows import dict_row

            pg_url = os.environ.get('DATABASE_URL', '').replace('postgresql+psycopg://', 'postgresql://')
            with psycopg.connect(pg_url, row_factory=dict_row) as conn:
                with conn.cursor() as cur:
                    cur.execute("""
                        SELECT vr.*, v.name as vehicle_name, v.registration as vehicle_registration
                        FROM vehicle_reservations vr
                        LEFT JOIN vehicles v ON vr.vehicle_id = v.id
                        ORDER BY vr.start_date DESC
                    """)
                    for row in cur.fetchall():
                        start = row['start_date']
                        end = row['end_date']
                        if end and end < today:
                            status = 'Завршена'
                        elif start and start <= today:
                            status = 'У току'
                        else:
                            status = 'Планирана'
                        activities.append({**row, 'computed_status': status})
        except Exception as exc:
            logger.warning("Failed to load activities: %s", exc)

    stats = {
        'completed': sum(1 for a in activities if a['computed_status'] == 'Завршена'),
        'active': sum(1 for a in activities if a['computed_status'] == 'У току'),
        'planned': sum(1 for a in activities if a['computed_status'] == 'Планирана'),
    }

    return render_template(
        'terenska_aktivnost.html',
        page_title='Теренска активност',
        activities=activities,
        stats=stats,
        today=today,
    )


def render_business_trip_request_page(*, get_museum_vehicles):
    vehicles = get_museum_vehicles()
    active_vehicles = [vehicle for vehicle in vehicles if vehicle.get('status') == 'Активно']
    department_heads = [
        {'name': 'др Биљана Митровић', 'role': 'шеф геолошког одељења'},
        {'name': 'Верица Стојановић', 'role': 'шеф биолошког одељења'},
        {'name': 'Душица Ивић', 'role': 'помоћник директора-руководилац одељења финансија'},
    ]
    return render_template(
        'zahtevi/zahtev_sluzbeni_put.html',
        page_title='Захтев за службени пут',
        user_full_name=session.get('user_name', ''),
        vehicles=active_vehicles,
        department_heads=department_heads,
    )


def api_field_trip_create(*, get_vehicle_reservations, save_reservations):
    """Create field trip request with vehicle reservation and timesheet entries."""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'message': 'Нема података'})

        result = {
            'success': True,
            'vehicle_reserved': False,
            'timesheet_updated': False,
            'message': 'Захтев за службени пут је креиран',
        }

        vehicle_id = data.get('vehicle_id')
        if vehicle_id and str(vehicle_id) != 'sopstveni':
            try:
                if os.environ.get('DATABASE_URL'):
                    import psycopg
                    from psycopg.rows import dict_row

                    pg_url = os.environ.get('DATABASE_URL', '').replace('postgresql+psycopg://', 'postgresql://')
                    with psycopg.connect(pg_url, row_factory=dict_row) as conn:
                        with conn.cursor() as cur:
                            cur.execute(
                                """
                                INSERT INTO vehicle_reservations (
                                    vehicle_id, reserved_by, purpose, start_date, end_date,
                                    destination, notes, status
                                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                                """,
                                (
                                    int(vehicle_id),
                                    session.get('user_name', session.get('user_email', 'system')),
                                    f"Теренски рад: {data.get('purpose', '')}",
                                    data.get('start_date'),
                                    data.get('end_date'),
                                    data.get('location', ''),
                                    f"Службени пут - {data.get('location')}",
                                    'Активна',
                                ),
                            )
                            conn.commit()
                            result['vehicle_reserved'] = True
                else:
                    reservations = get_vehicle_reservations()
                    existing_ids = [
                        res.get('id') for res in reservations
                        if isinstance(res.get('id'), int)
                    ]
                    next_id = (max(existing_ids) + 1) if existing_ids else 1
                    reservations.append(
                        {
                            'id': next_id,
                            'vehicle_id': int(vehicle_id),
                            'employee_name': session.get('user_name', ''),
                            'start_date': data.get('start_date'),
                            'end_date': data.get('end_date'),
                            'purpose': f"Теренски рад: {data.get('purpose', '')}",
                            'destination': data.get('location', ''),
                            'created_by': session.get('user_email', 'system'),
                            'created_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                        }
                    )
                    save_reservations()
                    result['vehicle_reserved'] = True
            except Exception as exc:
                logger.error("Vehicle reservation error: %s", exc)
                result['vehicle_error'] = str(exc)

        if data.get('update_timesheet') and data.get('start_date') and data.get('end_date'):
            try:
                start = datetime.strptime(data['start_date'], '%Y-%m-%d')
                end = datetime.strptime(data['end_date'], '%Y-%m-%d')

                if os.environ.get('DATABASE_URL'):
                    import psycopg
                    from psycopg.rows import dict_row

                    pg_url = os.environ.get('DATABASE_URL', '').replace('postgresql+psycopg://', 'postgresql://')
                    user_name = session.get('user_name', '')

                    with psycopg.connect(pg_url, row_factory=dict_row) as conn:
                        with conn.cursor() as cur:
                            days_recorded = 0
                            current = start
                            while current <= end:
                                cur.execute(
                                    """
                                    SELECT id FROM timesheet_reports
                                    WHERE employee_name = %s AND month = %s AND year = %s
                                    """,
                                    (user_name, current.month, current.year),
                                )
                                report = cur.fetchone()
                                if report:
                                    cur.execute(
                                        """
                                        INSERT INTO timesheet_report_days (report_id, day, work_outside)
                                        VALUES (%s, %s, 8)
                                        ON CONFLICT (report_id, day)
                                        DO UPDATE SET work_outside = 8, work_in_museum = 0
                                        """,
                                        (report['id'], current.day),
                                    )
                                    days_recorded += 1
                                current += timedelta(days=1)
                            conn.commit()
                            if days_recorded > 0:
                                result['timesheet_updated'] = True
                            else:
                                result['timesheet_skipped'] = True
            except Exception as exc:
                logger.error("Timesheet update error: %s", exc)
                result['timesheet_error'] = str(exc)

        return jsonify(result)
    except Exception as exc:
        logger.error("Field trip creation error: %s", exc)
        return jsonify({'success': False, 'message': str(exc)})


def _find_destination(dest_input):
    dest_upper = dest_input.upper().strip()
    for dest_name, info in DESTINATION_COORDS.items():
        if dest_name in dest_upper or dest_upper in dest_name:
            return dest_name, info
    for dest_name, info in DESTINATION_COORDS.items():
        for dest_word in dest_name.split():
            for input_word in dest_upper.split():
                if dest_word in input_word or input_word in dest_word:
                    return dest_name, info
    return None, None


def api_route_calculate(*, get_museum_vehicles):
    """Calculate route distance and toll from static database."""
    try:
        data = request.get_json()
        destinations = data.get('destinations', [])
        if not destinations:
            single_dest = data.get('destination', '').strip()
            if single_dest:
                destinations = [single_dest]
        if not destinations:
            return jsonify({'success': False, 'message': 'Није унета дестинација'})

        vehicle_id = data.get('vehicle_id')
        fuel_consumption = 10.0
        fuel_type = 'Бензин'
        vehicle_name = ''

        if vehicle_id == 'sopstveni':
            fuel_consumption = float(data.get('consumption', 8.0))
            fuel_type = data.get('fuel_type', 'BMB 95')
            vehicle_name = 'Сопствено возило'
        elif vehicle_id:
            for vehicle in get_museum_vehicles():
                if vehicle['id'] == int(vehicle_id):
                    fuel_consumption = vehicle.get('fuel_consumption', 10.0)
                    fuel_type = vehicle.get('fuel_type', 'Бензин')
                    vehicle_name = vehicle.get('name', '')
                    break

        fuel_price = FUEL_PRICES_LOCAL.get(fuel_type, 174)
        total_distance = 0
        total_toll = 0
        matched_destinations = []
        toll_stations = set()
        route_segments = []
        previous_location = 'Београд'

        # Check if OSM route geometry was provided for toll matching
        route_geometry = data.get('route_geometry')
        geometry_toll = 0
        geometry_toll_stations = set()
        if route_geometry and isinstance(route_geometry, list) and len(route_geometry) > 10:
            geometry_toll, geometry_toll_stations, _ = _toll_from_geometry(route_geometry)

        for dest_input in destinations:
            matched_dest, dest_info = _find_destination(dest_input)
            if dest_info:
                segment_km = dest_info['km']
                toll_station = dest_info.get('toll_station')
                segment_toll = 0
                if toll_station and not geometry_toll_stations:
                    toll_one_way = TOLL_PRICES_FROM_BELGRADE.get(toll_station, 0)
                    if toll_station not in toll_stations:
                        total_toll += toll_one_way
                        segment_toll = toll_one_way
                        toll_stations.add(toll_station)
                total_distance += segment_km
                matched_destinations.append(matched_dest)
                route_segments.append(
                    {
                        'from': previous_location,
                        'to': matched_dest,
                        'distance_km': segment_km,
                        'toll_station': toll_station if not geometry_toll_stations else None,
                        'toll_cost': segment_toll,
                        'cumulative_km': total_distance,
                    }
                )
                previous_location = matched_dest
            else:
                matched_destinations.append(dest_input)
                route_segments.append(
                    {
                        'from': previous_location,
                        'to': dest_input,
                        'distance_km': 0,
                        'toll_station': None,
                        'toll_cost': 0,
                        'cumulative_km': total_distance,
                        'unknown': True,
                    }
                )
                previous_location = dest_input

        # When geometry is available, use geometry-based tolls
        if geometry_toll_stations:
            total_toll = geometry_toll
            toll_stations = geometry_toll_stations
            return_toll = geometry_toll  # same toll on return
            # Assign toll info to first segment for display
            if route_segments:
                route_segments[0]['toll_station'] = ', '.join(sorted(toll_stations))
                route_segments[0]['toll_cost'] = total_toll
        else:
            return_toll = total_toll
            if matched_destinations:
                last_dest = matched_destinations[-1].upper()
                for dest_name, info in DESTINATION_COORDS.items():
                    if dest_name in last_dest or last_dest in dest_name:
                        toll_station = info.get('toll_station')
                        return_toll = TOLL_PRICES_FROM_BELGRADE.get(toll_station, 0) if toll_station else 0
                        break

        return_distance = total_distance
        if matched_destinations:
            last_dest = matched_destinations[-1].upper()
            for dest_name, info in DESTINATION_COORDS.items():
                if dest_name in last_dest or last_dest in dest_name:
                    return_distance = info['km']
                    break

        route_segments.append(
            {
                'from': previous_location,
                'to': 'Београд',
                'distance_km': return_distance,
                'toll_station': 'повратак',
                'toll_cost': return_toll,
                'is_return': True,
            }
        )

        if total_distance <= 0:
            return jsonify(
                {
                    'success': False,
                    'destination': ', '.join(destinations),
                    'distance_km': 0,
                    'message': 'Дестинација није пронађена. Унесите километражу ручно.',
                    'known_destinations': list(DESTINATION_COORDS.keys())[:20],
                }
            )

        distance_round_trip = total_distance + return_distance
        toll_round_trip = total_toll + return_toll
        fuel_liters = distance_round_trip * fuel_consumption / 100
        fuel_cost = int(fuel_liters * fuel_price)
        route_parts = ['Београд'] + matched_destinations + ['Београд']
        return jsonify(
            {
                'success': True,
                'source': 'Database',
                'destinations': matched_destinations,
                'destination': ' → '.join(matched_destinations),
                'distance_km_outbound': total_distance,
                'distance_km_return': return_distance,
                'distance_round_trip': distance_round_trip,
                'toll_station': ', '.join(sorted(toll_stations)) if toll_stations else 'без путарине',
                'toll_stations': list(toll_stations),
                'toll_outbound': total_toll,
                'toll_return': return_toll,
                'toll_round_trip': toll_round_trip,
                'fuel_consumption': fuel_consumption,
                'fuel_type': fuel_type,
                'fuel_liters': round(fuel_liters, 1),
                'fuel_price_per_liter': fuel_price,
                'fuel_cost': fuel_cost,
                'vehicle_name': vehicle_name,
                'total_transport_cost': fuel_cost + toll_round_trip,
                'route_info': ' → '.join(route_parts),
                'segments': route_segments,
            }
        )
    except Exception as exc:
        logger.error("Route calculation error: %s", exc)
        return jsonify({'success': False, 'message': str(exc)})


_CYR_TO_LAT = str.maketrans({
    'А': 'A', 'Б': 'B', 'В': 'V', 'Г': 'G', 'Д': 'D', 'Ђ': 'Đ',
    'Е': 'E', 'Ж': 'Ž', 'З': 'Z', 'И': 'I', 'Ј': 'J', 'К': 'K',
    'Л': 'L', 'Љ': 'Lj', 'М': 'M', 'Н': 'N', 'Њ': 'Nj', 'О': 'O',
    'П': 'P', 'Р': 'R', 'С': 'S', 'Т': 'T', 'Ћ': 'Ć', 'У': 'U',
    'Ф': 'F', 'Х': 'H', 'Ц': 'C', 'Ч': 'Č', 'Џ': 'Dž', 'Ш': 'Š',
    'а': 'a', 'б': 'b', 'в': 'v', 'г': 'g', 'д': 'd', 'ђ': 'đ',
    'е': 'e', 'ж': 'ž', 'з': 'z', 'и': 'i', 'ј': 'j', 'к': 'k',
    'л': 'l', 'љ': 'lj', 'м': 'm', 'н': 'n', 'њ': 'nj', 'о': 'o',
    'п': 'p', 'р': 'r', 'с': 's', 'т': 't', 'ћ': 'ć', 'у': 'u',
    'ф': 'f', 'х': 'h', 'ц': 'c', 'ч': 'č', 'џ': 'dž', 'ш': 'š',
})


def _cyr_to_lat(text):
    """Transliterate Serbian Cyrillic to Latin."""
    return text.translate(_CYR_TO_LAT)


def _nominatim_search(query):
    """Search Nominatim with Cyrillic→Latin fallback on failure."""
    headers = {
        'Accept': 'application/json',
        'Accept-Language': 'sr,en',
        'User-Agent': 'MuseumInfoSystem/1.0',
    }
    params = {'format': 'jsonv2', 'limit': 1, 'q': query}

    for attempt in range(2):
        try:
            response = http_requests.get(
                'https://nominatim.openstreetmap.org/search',
                params=params, headers=headers, timeout=15,
            )
            if response.status_code == 429:
                # Rate limited — wait and retry once
                import time
                time.sleep(1.5)
                continue
            response.raise_for_status()
            results = response.json()
            if results:
                return results[0]
            break  # 200 but no results — don't retry, fall through to Latin
        except Exception as exc:
            logger.warning("Nominatim search failed for %r: %s", query, exc)
            break

    # Fallback: transliterate Cyrillic to Latin and retry
    latin_query = _cyr_to_lat(query)
    if latin_query != query:
        try:
            import time
            time.sleep(1.1)  # respect Nominatim rate limit
            response = http_requests.get(
                'https://nominatim.openstreetmap.org/search',
                params={**params, 'q': latin_query}, headers=headers, timeout=15,
            )
            response.raise_for_status()
            results = response.json()
            if results:
                return results[0]
        except Exception as exc:
            logger.warning("Nominatim Latin fallback failed for %r: %s", latin_query, exc)

    return None


# Fixed coordinates for the museum — avoids geocoding a known location
_MUSEUM_POINT = {
    'lat': 44.81565,
    'lon': 20.46175,
    'display_name': 'Природњачки музеј, Београд',
}


def api_osm_route_preview():
    """Resolve OSM locations and build a route preview through server-side requests."""
    try:
        data = request.get_json() or {}
        locations = [str(location).strip() for location in data.get('locations', []) if str(location).strip()]
        if not locations:
            return jsonify({'success': False, 'message': 'Није унета ниједна локација'}), 400

        # Start and end at museum (known coordinates, no geocoding needed)
        resolved_points = [dict(_MUSEUM_POINT)]

        for location in locations:
            match = _nominatim_search(f'{location}, Serbia')
            if not match:
                return jsonify({'success': False, 'message': f'Локација није пронађена: {location}'}), 404

            resolved_points.append(
                {
                    'lat': float(match['lat']),
                    'lon': float(match['lon']),
                    'display_name': match.get('display_name', location),
                }
            )
            import time
            time.sleep(1.1)  # respect Nominatim 1 req/sec limit

        resolved_points.append(dict(_MUSEUM_POINT))

        coordinates = ';'.join(f"{point['lon']},{point['lat']}" for point in resolved_points)
        route_response = http_requests.get(
            f'https://router.project-osrm.org/route/v1/driving/{coordinates}',
            params={
                'overview': 'full',
                'geometries': 'geojson',
                'steps': 'true',
            },
            headers={
                'Accept': 'application/json',
                'User-Agent': 'MuseumInfoSystem/1.0',
            },
            timeout=20,
        )
        route_response.raise_for_status()
        route_data = route_response.json()
        if not route_data.get('routes'):
            return jsonify({'success': False, 'message': 'OSM рута није пронађена'}), 404

        route = route_data['routes'][0]
        return jsonify(
            {
                'success': True,
                'resolved_points': resolved_points,
                'route': {
                    'distance': route.get('distance', 0),
                    'duration': route.get('duration', 0),
                    'geometry': route.get('geometry', {}),
                    'legs': [
                        {
                            'distance': leg.get('distance', 0),
                            'duration': leg.get('duration', 0),
                        }
                        for leg in route.get('legs', [])
                    ],
                },
            }
        )
    except Exception as exc:
        logger.error("OSM route preview error: %s", exc)
        return jsonify({'success': False, 'message': str(exc)}), 500


def _build_procurement_document(data):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('ДИРЕКТОРУ ПРИРОДЊАЧКОГ МУЗЕЈА, БЕОГРАД')
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    header_table = doc.add_table(rows=2, cols=3)
    header_table.style = 'Table Grid'
    header_table.cell(0, 0).merge(header_table.cell(1, 0))
    header_table.cell(0, 0).text = '[LOGO]'
    header_table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_table.cell(0, 1).merge(header_table.cell(1, 1))
    cell = header_table.cell(0, 1)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run('ЗАХТЕВ\n')
    run.bold = True
    run.font.size = Pt(16)
    paragraph.add_run('за набавку / наруџбу')
    header_table.cell(0, 2).text = f"Датум подношења:\n{data.get('datum', '')}"

    doc.add_paragraph()
    applicant_table = doc.add_table(rows=1, cols=2)
    applicant_table.style = 'Table Grid'
    applicant_table.cell(0, 0).text = f"Подносилац захтева:\n{data.get('podnosilac', '')}"
    applicant_table.cell(0, 1).text = "Потпис:\n\n"
    doc.add_paragraph()

    items = data.get('items', [])
    items_table = doc.add_table(rows=len(items) + 3, cols=5)
    items_table.style = 'Table Grid'
    items_table.cell(0, 0).merge(items_table.cell(1, 0))
    items_table.cell(0, 0).text = 'Р.бр.'
    items_table.cell(0, 1).merge(items_table.cell(1, 1))
    items_table.cell(0, 1).text = 'Добро / услуга'
    items_table.cell(0, 2).merge(items_table.cell(0, 3))
    items_table.cell(0, 2).text = 'Динара'
    items_table.cell(0, 4).merge(items_table.cell(1, 4))
    items_table.cell(0, 4).text = 'Примио добро / услугу / средства:'
    items_table.cell(1, 2).text = 'Процењено'
    items_table.cell(1, 3).text = 'Реализовано'

    for idx, item in enumerate(items):
        row_idx = idx + 2
        items_table.cell(row_idx, 0).text = f"{item.get('rbr', idx + 1)}."
        items_table.cell(row_idx, 1).text = item.get('description', '')
        items_table.cell(row_idx, 2).text = f"{item.get('estimated', 0):,.2f}"
        items_table.cell(row_idx, 3).text = f"{item.get('realized', 0):,.2f}"
        items_table.cell(row_idx, 4).text = item.get('receiver', '')

    total_row = len(items) + 2
    items_table.cell(total_row, 0).merge(items_table.cell(total_row, 1))
    items_table.cell(total_row, 0).text = 'СВЕГА:'
    items_table.cell(total_row, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    items_table.cell(total_row, 2).text = f"{data.get('totalEstimated', 0):,.2f}"
    items_table.cell(total_row, 3).text = f"{data.get('totalRealized', 0):,.2f}"

    doc.add_paragraph()
    approval_table = doc.add_table(rows=2, cols=2)
    approval_table.style = 'Table Grid'
    approval_table.cell(0, 0).text = (
        f"На терет активности:\n{data.get('teretAktivnosti', '')}\n{data.get('teretAktivnostiOpis', '')}"
    )
    approval_table.cell(0, 1).text = f"Сагласан руководилац:\n{data.get('saglasanRukovodilac', '')}\n\n"
    approval_table.cell(1, 0).text = (
        f"Да постоје расположива наменска средства, потврђује шеф рачуноводства:\n{data.get('sefRacunovodstva', '')}\n\n"
    )
    approval_table.cell(1, 1).text = f"Одобрава набавку / наруџбу, директор:\n{data.get('direktor', '')}\n\n"
    return doc


def _send_docx_document(doc, download_name):
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return send_file(
        file_stream,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=download_name,
    )


def api_nabavka_save():
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'message': 'Нема података'})

        import psycopg
        from psycopg.rows import dict_row

        db_url = os.environ.get('DATABASE_URL', '').replace('postgresql+psycopg://', 'postgresql://')
        with psycopg.connect(conninfo=db_url, row_factory=dict_row) as conn:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    CREATE TABLE IF NOT EXISTS procurement_requests (
                        id SERIAL PRIMARY KEY,
                        datum DATE NOT NULL,
                        podnosilac VARCHAR(255) NOT NULL,
                        items JSONB NOT NULL,
                        total_estimated DECIMAL(12,2) DEFAULT 0,
                        total_realized DECIMAL(12,2) DEFAULT 0,
                        teret_aktivnosti VARCHAR(100),
                        teret_aktivnosti_opis TEXT,
                        saglasan_rukovodilac VARCHAR(255),
                        sef_racunovodstva VARCHAR(255),
                        direktor VARCHAR(255),
                        status VARCHAR(50) DEFAULT 'pending',
                        user_email VARCHAR(255),
                        created_at TIMESTAMP DEFAULT NOW(),
                        updated_at TIMESTAMP DEFAULT NOW()
                    )
                    """
                )
                cur.execute(
                    """
                    INSERT INTO procurement_requests
                    (datum, podnosilac, items, total_estimated, total_realized,
                     teret_aktivnosti, teret_aktivnosti_opis, saglasan_rukovodilac,
                     sef_racunovodstva, direktor, user_email)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                    RETURNING id
                    """,
                    (
                        data.get('datum'),
                        data.get('podnosilac'),
                        json.dumps(data.get('items', []), ensure_ascii=False),
                        data.get('totalEstimated', 0),
                        data.get('totalRealized', 0),
                        data.get('teretAktivnosti'),
                        data.get('teretAktivnostiOpis'),
                        data.get('saglasanRukovodilac'),
                        data.get('sefRacunovodstva'),
                        data.get('direktor'),
                        session.get('user_email'),
                    ),
                )
                new_id = cur.fetchone()['id']
                conn.commit()
        return jsonify({'success': True, 'message': f'Захтев сачуван (ID: {new_id})', 'id': new_id})
    except Exception as exc:
        logger.exception("Error saving procurement request")
        return jsonify({'success': False, 'message': str(exc)})


def api_nabavka_list():
    try:
        import psycopg
        from psycopg.rows import dict_row

        db_url = os.environ.get('DATABASE_URL', '').replace('postgresql+psycopg://', 'postgresql://')
        with psycopg.connect(conninfo=db_url, row_factory=dict_row) as conn:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    SELECT EXISTS (
                        SELECT FROM information_schema.tables
                        WHERE table_name = 'procurement_requests'
                    )
                    """
                )
                if not cur.fetchone()['exists']:
                    return jsonify({'success': True, 'requests': []})

                user_email = session.get('user_email')
                is_admin = session.get('user_role') == 'admin' or session.get('is_admin', False)
                if is_admin:
                    cur.execute(
                        """
                        SELECT id, datum::text, podnosilac, items, total_estimated as procenjeno,
                               status, created_at
                        FROM procurement_requests
                        ORDER BY created_at DESC
                        LIMIT 50
                        """
                    )
                else:
                    cur.execute(
                        """
                        SELECT id, datum::text, podnosilac, items, total_estimated as procenjeno,
                               status, created_at
                        FROM procurement_requests
                        WHERE user_email = %s
                        ORDER BY created_at DESC
                        LIMIT 50
                        """,
                        (user_email,),
                    )
                requests = cur.fetchall()

        formatted = []
        for req in requests:
            items = req['items'] if isinstance(req['items'], list) else json.loads(req['items'])
            first = items[0] if isinstance(items, list) and items else None
            desc = first.get('description', '') if isinstance(first, dict) else ''
            opis = (str(desc or '')[:50] + '...') if first is not None else 'Без описа'
            formatted.append(
                {
                    'id': req['id'],
                    'datum': req['datum'],
                    'opis': opis,
                    'procenjeno': float(req['procenjeno'] or 0),
                    'status': req['status'],
                }
            )
        return jsonify({'success': True, 'requests': formatted})
    except Exception as exc:
        logger.error("Error listing procurement requests: %s", exc)
        return jsonify({'success': False, 'message': str(exc)}), 500


def api_nabavka_export_word():
    try:
        data = request.get_json()
        return _send_docx_document(
            _build_procurement_document(data),
            f"Zahtev_za_nabavku_{data.get('datum', 'unknown')}.docx",
        )
    except Exception as exc:
        logger.exception("Error exporting procurement request")
        return jsonify({'success': False, 'message': str(exc)}), 500


def api_nabavka_export_word_by_id(request_id, *, can_access_owned_record):
    try:
        import psycopg
        from psycopg.rows import dict_row

        db_url = os.environ.get('DATABASE_URL', '').replace('postgresql+psycopg://', 'postgresql://')
        with psycopg.connect(conninfo=db_url, row_factory=dict_row) as conn:
            with conn.cursor() as cur:
                cur.execute("SELECT * FROM procurement_requests WHERE id = %s", (request_id,))
                req = cur.fetchone()
                if not req:
                    return jsonify({'success': False, 'message': 'Захтев није пронађен'}), 404
                if not can_access_owned_record(
                    req.get('user_email'),
                    session.get('user_email', ''),
                    session.get('user_role', ''),
                ):
                    return jsonify({'success': False, 'message': 'Немате приступ овом захтеву'}), 403

        items = req['items'] if isinstance(req['items'], list) else json.loads(req['items'])
        data = {
            'datum': str(req['datum']),
            'podnosilac': req['podnosilac'],
            'items': items,
            'totalEstimated': float(req['total_estimated'] or 0),
            'totalRealized': float(req['total_realized'] or 0),
            'teretAktivnosti': req['teret_aktivnosti'] or '',
            'teretAktivnostiOpis': req['teret_aktivnosti_opis'] or '',
            'saglasanRukovodilac': req['saglasan_rukovodilac'] or '',
            'sefRacunovodstva': req['sef_racunovodstva'] or '',
            'direktor': req['direktor'] or '',
        }
        return _send_docx_document(
            _build_procurement_document(data),
            f"Zahtev_za_nabavku_{request_id}_{data['datum']}.docx",
        )
    except Exception as exc:
        logger.exception("Error exporting procurement request by id")
        return jsonify({'success': False, 'message': str(exc)}), 500


def _financial_plan_totals(payload):
    _selected_year, years_data = _normalize_financial_plan_payload(payload)
    totals_by_year = {
        str(year): _financial_plan_year_total(year_data)
        for year, year_data in years_data.items()
    }
    grand_total = sum(totals_by_year.values())
    return totals_by_year, grand_total


def api_finansijski_plan_save(*, get_postgres_connection):
    try:
        data = request.get_json()
        years_data = data.get('years', {})
        plan_payload = {
            'selectedYear': str(data.get('selectedYear') or '').strip() or None,
            'years': years_data,
        }
        totals_by_year, grand_total = _financial_plan_totals(plan_payload)
        raw_plan_id = data.get('id')
        plan_id = int(raw_plan_id) if str(raw_plan_id or '').strip().isdigit() else None

        with get_postgres_connection() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    CREATE TABLE IF NOT EXISTS financial_plans (
                        id SERIAL PRIMARY KEY,
                        odeljenje VARCHAR(100),
                        odeljenje_text VARCHAR(200),
                        kustos VARCHAR(200),
                        datum_izrade DATE,
                        plan_data JSONB,
                        total_2026 DECIMAL(15,2),
                        total_2027 DECIMAL(15,2),
                        total_2028 DECIMAL(15,2),
                        grand_total DECIMAL(15,2),
                        user_email VARCHAR(200),
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                    )
                    """
                )
                if plan_id is not None:
                    cur.execute(
                        "SELECT user_email FROM financial_plans WHERE id = %s",
                        (plan_id,),
                    )
                    owner_row = cur.fetchone()
                    if not owner_row:
                        return jsonify({'success': False, 'message': 'План није пронађен'}), 404

                    owner_email = owner_row[0]
                    current_email = session.get('user_email', '')
                    current_role = session.get('user_role', '')
                    if owner_email != current_email and current_role != 'admin':
                        return jsonify({'success': False, 'message': 'Немате приступ овом плану'}), 403

                    cur.execute(
                        """
                        UPDATE financial_plans
                        SET odeljenje = %s,
                            odeljenje_text = %s,
                            kustos = %s,
                            datum_izrade = %s,
                            plan_data = %s,
                            total_2026 = %s,
                            total_2027 = %s,
                            total_2028 = %s,
                            grand_total = %s
                        WHERE id = %s
                        RETURNING id
                        """,
                        (
                            data.get('odeljenje'),
                            data.get('odeljenjeText'),
                            data.get('kustos'),
                            data.get('datumIzrade'),
                            json.dumps(plan_payload),
                            totals_by_year.get('2026', 0),
                            totals_by_year.get('2027', 0),
                            totals_by_year.get('2028', 0),
                            grand_total,
                            plan_id,
                        ),
                    )
                    plan_id = cur.fetchone()[0]
                    updated = True
                else:
                    cur.execute(
                        """
                        INSERT INTO financial_plans
                        (odeljenje, odeljenje_text, kustos, datum_izrade, plan_data,
                         total_2026, total_2027, total_2028, grand_total, user_email)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                        RETURNING id
                        """,
                        (
                            data.get('odeljenje'),
                            data.get('odeljenjeText'),
                            data.get('kustos'),
                            data.get('datumIzrade'),
                            json.dumps(plan_payload),
                            totals_by_year.get('2026', 0),
                            totals_by_year.get('2027', 0),
                            totals_by_year.get('2028', 0),
                            grand_total,
                            session.get('user_email', ''),
                        ),
                    )
                    plan_id = cur.fetchone()[0]
                    updated = False
                conn.commit()
        return jsonify({'success': True, 'id': plan_id, 'updated': updated})
    except Exception as exc:
        logger.error("Error saving financial plan: %s", exc)
        return jsonify({'success': False, 'message': str(exc)}), 500


def api_finansijski_plan_list(*, get_postgres_connection, current_user_is_admin):
    try:
        with get_postgres_connection() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    SELECT EXISTS (
                        SELECT FROM information_schema.tables
                        WHERE table_name = 'financial_plans'
                    )
                    """
                )
                if not cur.fetchone()[0]:
                    return jsonify({'success': True, 'plans': []})

                if current_user_is_admin():
                    cur.execute(
                        """
                        SELECT id, datum_izrade, odeljenje_text, kustos, grand_total,
                               user_email, total_2026, total_2027, total_2028, created_at
                        FROM financial_plans
                        ORDER BY created_at DESC
                        """
                    )
                else:
                    cur.execute(
                        """
                        SELECT id, datum_izrade, odeljenje_text, kustos, grand_total,
                               user_email, total_2026, total_2027, total_2028, created_at
                        FROM financial_plans
                        WHERE user_email = %s
                        ORDER BY created_at DESC
                        """,
                        (session.get('user_email', ''),),
                    )
                rows = cur.fetchall()

        plans = []
        for row in rows:
            plans.append(
                {
                    'id': row[0],
                    'datum': str(row[1]) if row[1] else '',
                    'odeljenje': row[2] or '',
                    'kustos': row[3] or '',
                    'ukupno': float(row[4]) if row[4] else 0,
                    'user_email': row[5] or '',
                    'total_2026': float(row[6]) if row[6] else 0,
                    'total_2027': float(row[7]) if row[7] else 0,
                    'total_2028': float(row[8]) if row[8] else 0,
                    'created_at': row[9].isoformat() if row[9] else None,
                }
            )
        return jsonify({'success': True, 'plans': plans})
    except Exception as exc:
        logger.error("Error listing financial plans: %s", exc)
        return jsonify({'success': False, 'message': str(exc)}), 500


def api_finansijski_plan_get(
    plan_id,
    *,
    get_postgres_connection,
    can_access_owned_record,
):
    try:
        with get_postgres_connection() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    SELECT odeljenje, odeljenje_text, kustos, datum_izrade, plan_data, user_email
                    FROM financial_plans
                    WHERE id = %s
                    """,
                    (plan_id,),
                )
                row = cur.fetchone()
                if not row:
                    return jsonify({'success': False, 'message': 'План није пронађен'}), 404

                owner_email = row[5]
                if not can_access_owned_record(
                    owner_email,
                    session.get('user_email', ''),
                    session.get('user_role', ''),
                ):
                    return jsonify({'success': False, 'message': 'Немате приступ овом плану'}), 403

        selected_year, years_data = _normalize_financial_plan_payload(row[4] if row[4] else {})
        if not selected_year:
            available_years = sorted(years_data.keys(), key=lambda value: int(value)) if years_data else []
            selected_year = available_years[0] if available_years else None

        return jsonify(
            {
                'success': True,
                'plan': {
                    'id': plan_id,
                    'odeljenje': row[0] or '',
                    'odeljenjeText': row[1] or '',
                    'kustos': row[2] or '',
                    'imePrezime': row[2] or '',
                    'datumIzrade': str(row[3]) if row[3] else '',
                    'selectedYear': selected_year,
                    'years': years_data,
                },
            }
        )
    except Exception as exc:
        logger.exception("Error loading financial plan by id")
        return jsonify({'success': False, 'message': str(exc)}), 500


def _build_financial_plan_document(data):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    selected_year, years_data = _normalize_financial_plan_payload(data)
    if not years_data and isinstance(data, dict):
        selected_year, years_data = _normalize_financial_plan_payload(data.get('years', {}))

    if selected_year and selected_year in years_data:
        years_to_render = [selected_year]
    else:
        years_to_render = sorted(years_data.keys(), key=lambda value: int(value)) if years_data else []

    for index, year in enumerate(years_to_render):
        year_data = years_data.get(year, {})
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run("Предлог финансијског плана")
        run.bold = True
        run.font.size = Pt(14)

        year_para = doc.add_paragraph()
        year_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = year_para.add_run(f"Година: {year}")
        run.bold = True
        run.font.size = Pt(12)

        dept_para = doc.add_paragraph()
        dept_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dept_para.add_run(f"Одељење: {data.get('odeljenjeText', '')}").font.size = Pt(11)

        doc.add_paragraph()
        year_total = 0
        for category_key, category_name in FINANCIAL_PLAN_CATEGORY_NAMES.items():
            items = year_data.get(category_key, [])
            if not items:
                continue

            cat_header = doc.add_paragraph()
            run = cat_header.add_run(category_name)
            run.bold = True
            run.font.size = Pt(11)

            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'р.б.'
            hdr_cells[1].text = 'Активност'
            hdr_cells[2].text = 'Износ'
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            category_total = 0
            for item in items:
                row_cells = table.add_row().cells
                row_cells[0].text = f"{item.get('rbr', '')}."
                row_cells[1].text = item.get('activity', '')
                amount = item.get('amount', 0)
                row_cells[2].text = f"{amount:,.0f}".replace(',', '.')
                row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                category_total += amount

            total_row = table.add_row().cells
            total_row[0].text = ''
            total_row[1].text = 'УКУПНО:'
            total_row[1].paragraphs[0].runs[0].bold = True
            total_row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            total_row[2].text = f"{category_total:,.0f}".replace(',', '.')
            total_row[2].paragraphs[0].runs[0].bold = True
            total_row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            year_total += category_total
            doc.add_paragraph()

        year_total_para = doc.add_paragraph()
        year_total_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = year_total_para.add_run(f"УКУПНО ЗА {year}. ГОДИНУ: {year_total:,.0f} РСД".replace(',', '.'))
        run.bold = True
        run.font.size = Pt(12)
        if index < len(years_to_render) - 1:
            doc.add_page_break()

    doc.add_paragraph()
    doc.add_paragraph()
    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_para.add_run(f"Кустос: {data.get('kustos', '')}")
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_para.add_run(f"Датум: {data.get('datumIzrade', '')}")
    return doc


def api_finansijski_plan_export_word():
    try:
        data = request.get_json()
        return _send_docx_document(
            _build_financial_plan_document(data),
            f"Finansijski_plan_{data.get('datumIzrade', 'unknown')}.docx",
        )
    except Exception as exc:
        logger.exception("Error exporting financial plan")
        return jsonify({'success': False, 'message': str(exc)}), 500


def api_finansijski_plan_export_word_by_id(
    plan_id,
    *,
    get_postgres_connection,
    can_access_owned_record,
):
    try:
        with get_postgres_connection() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    SELECT odeljenje, odeljenje_text, kustos, datum_izrade, plan_data
                    FROM financial_plans WHERE id = %s
                    """,
                    (plan_id,),
                )
                row = cur.fetchone()
                if not row:
                    return jsonify({'success': False, 'message': 'План није пронађен'}), 404
                cur.execute("SELECT user_email FROM financial_plans WHERE id = %s", (plan_id,))
                owner_row = cur.fetchone()
                owner_email = owner_row[0] if owner_row else None
                if not can_access_owned_record(
                    owner_email,
                    session.get('user_email', ''),
                    session.get('user_role', ''),
                ):
                    return jsonify({'success': False, 'message': 'Немате приступ овом плану'}), 403

        data = {
            'odeljenje': row[0],
            'odeljenjeText': row[1],
            'kustos': row[2],
            'datumIzrade': str(row[3]) if row[3] else '',
            'years': row[4] if row[4] else {},
        }
        return _send_docx_document(
            _build_financial_plan_document(data),
            f"Finansijski_plan_{plan_id}_{data['datumIzrade']}.docx",
        )
    except Exception as exc:
        logger.exception("Error exporting financial plan by id")
        return jsonify({'success': False, 'message': str(exc)}), 500

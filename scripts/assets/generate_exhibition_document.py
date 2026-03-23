#!/usr/bin/env python3
"""
Generate Word document for permanent exhibition specimens
Output in Serbian language
"""

import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Load selected specimens
with open('exhibition_specimens.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

balkan_specimens = data['balkan_specimens']
world_specimens = data['world_specimens']

# Create document
doc = Document()

# Set up document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title
title = doc.add_heading('СТАЛНА ПОСТАВКА МИНЕРАЛОШКЕ ЗБИРКЕ', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Природњачки музеј у Београду')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(14)
subtitle_format.font.bold = True

# Date
date_para = doc.add_paragraph(f'Датум: {datetime.now().strftime("%d.%m.%Y.")}')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_format = date_para.runs[0]
date_format.font.size = Pt(11)
date_format.font.italic = True

doc.add_paragraph()  # spacing

# Introduction
intro = doc.add_paragraph()
intro.add_run('Овај документ садржи предлог минералошких примерака за сталну поставку музеја. ')
intro.add_run('Избор је направљен са циљем представљања репрезентативне минералошке разноврсности ')
intro.add_run('са посебним нагласком на примерке са подручја Балканског полуострва и бивше Југославије, ')
intro.add_run('као и значајне примерке из света.')

doc.add_paragraph()

# Summary statistics
stats = doc.add_paragraph()
stats.add_run('Укупно примерака: ').bold = True
stats.add_run(f'{len(balkan_specimens) + len(world_specimens)}\n')
stats.add_run('Примерци са Балкана: ').bold = True
stats.add_run(f'{len(balkan_specimens)}\n')
stats.add_run('Примерци из света: ').bold = True
stats.add_run(f'{len(world_specimens)}')

doc.add_page_break()

# ==================== BALKAN SPECIMENS ====================
heading1 = doc.add_heading('I. МИНЕРАЛИ СА БАЛКАНСКОГ ПОЛУОСТРВА И БИВШЕ ЈУГОСЛАВИЈЕ', 1)
heading1_format = heading1.runs[0]
heading1_format.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph(f'Укупно примерака: {len(balkan_specimens)}')
doc.add_paragraph()

for i, specimen in enumerate(balkan_specimens, 1):
    # Number and mineral name
    p = doc.add_paragraph()
    number_run = p.add_run(f'{i}. ')
    number_run.bold = True
    number_run.font.size = Pt(12)

    name_run = p.add_run(specimen['naziv'].upper())
    name_run.bold = True
    name_run.font.size = Pt(12)
    name_run.font.color.rgb = RGBColor(0, 0, 0)

    # Inventory number
    if specimen['inv_broj']:
        inv_para = doc.add_paragraph()
        inv_para.paragraph_format.left_indent = Inches(0.3)
        inv_run = inv_para.add_run(f"Инвентарни број: M{specimen['inv_broj']}")
        inv_run.font.size = Pt(11)
        inv_run.italic = True

    # Locality
    loc_para = doc.add_paragraph()
    loc_para.paragraph_format.left_indent = Inches(0.3)
    loc_label = loc_para.add_run('Локалитет: ')
    loc_label.bold = True
    loc_label.font.size = Pt(11)
    loc_value = loc_para.add_run(specimen['lokalitet'])
    loc_value.font.size = Pt(11)

    # Predmet (object type)
    if specimen['predmet']:
        obj_para = doc.add_paragraph()
        obj_para.paragraph_format.left_indent = Inches(0.3)
        obj_label = obj_para.add_run('Предмет: ')
        obj_label.bold = True
        obj_label.font.size = Pt(10)
        obj_value = obj_para.add_run(specimen['predmet'])
        obj_value.font.size = Pt(10)

    # Description
    if specimen['opis']:
        desc_para = doc.add_paragraph()
        desc_para.paragraph_format.left_indent = Inches(0.3)
        desc_label = desc_para.add_run('Опис: ')
        desc_label.bold = True
        desc_label.font.size = Pt(10)
        desc_value = desc_para.add_run(specimen['opis'])
        desc_value.font.size = Pt(10)

    # Comment
    if specimen['komentar']:
        comm_para = doc.add_paragraph()
        comm_para.paragraph_format.left_indent = Inches(0.3)
        comm_label = comm_para.add_run('Напомена: ')
        comm_label.bold = True
        comm_label.font.size = Pt(10)
        comm_value = comm_para.add_run(specimen['komentar'])
        comm_value.font.size = Pt(10)

    # Add spacing after each entry
    doc.add_paragraph()

doc.add_page_break()

# ==================== WORLD SPECIMENS ====================
heading2 = doc.add_heading('II. МИНЕРАЛИ ИЗ СВЕТА', 1)
heading2_format = heading2.runs[0]
heading2_format.font.color.rgb = RGBColor(102, 51, 0)

doc.add_paragraph(f'Укупно примерака: {len(world_specimens)}')
doc.add_paragraph()

for i, specimen in enumerate(world_specimens, 1):
    # Number and mineral name
    p = doc.add_paragraph()
    number_run = p.add_run(f'{i}. ')
    number_run.bold = True
    number_run.font.size = Pt(12)

    name_run = p.add_run(specimen['naziv'].upper())
    name_run.bold = True
    name_run.font.size = Pt(12)
    name_run.font.color.rgb = RGBColor(0, 0, 0)

    # Inventory number
    if specimen['inv_broj']:
        inv_para = doc.add_paragraph()
        inv_para.paragraph_format.left_indent = Inches(0.3)
        inv_run = inv_para.add_run(f"Инвентарни број: M{specimen['inv_broj']}")
        inv_run.font.size = Pt(11)
        inv_run.italic = True

    # Locality
    loc_para = doc.add_paragraph()
    loc_para.paragraph_format.left_indent = Inches(0.3)
    loc_label = loc_para.add_run('Локалитет: ')
    loc_label.bold = True
    loc_label.font.size = Pt(11)
    loc_value = loc_para.add_run(specimen['lokalitet'])
    loc_value.font.size = Pt(11)

    # Predmet (object type)
    if specimen['predmet']:
        obj_para = doc.add_paragraph()
        obj_para.paragraph_format.left_indent = Inches(0.3)
        obj_label = obj_para.add_run('Предмет: ')
        obj_label.bold = True
        obj_label.font.size = Pt(10)
        obj_value = obj_para.add_run(specimen['predmet'])
        obj_value.font.size = Pt(10)

    # Description
    if specimen['opis']:
        desc_para = doc.add_paragraph()
        desc_para.paragraph_format.left_indent = Inches(0.3)
        desc_label = desc_para.add_run('Опис: ')
        desc_label.bold = True
        desc_label.font.size = Pt(10)
        desc_value = desc_para.add_run(specimen['opis'])
        desc_value.font.size = Pt(10)

    # Comment
    if specimen['komentar']:
        comm_para = doc.add_paragraph()
        comm_para.paragraph_format.left_indent = Inches(0.3)
        comm_label = comm_para.add_run('Напомена: ')
        comm_label.bold = True
        comm_label.font.size = Pt(10)
        comm_value = comm_para.add_run(specimen['komentar'])
        comm_value.font.size = Pt(10)

    # Add spacing after each entry
    doc.add_paragraph()

# Add footer with notes
doc.add_page_break()
footer_heading = doc.add_heading('НАПОМЕНЕ', 2)
notes = doc.add_paragraph()
notes.add_run('• ').bold = True
notes.add_run('Овај избор представља предлог за сталну поставку заснован на постојећој минералошкој збирци музеја.\n')
notes.add_run('• ').bold = True
notes.add_run('Примерци су изабрани на основу репрезентативности, порекла и очуваности.\n')
notes.add_run('• ').bold = True
notes.add_run('Приоритет је дат примерцима са подручја Балкана и бивше Југославије.\n')
notes.add_run('• ').bold = True
notes.add_run('Сви примерци су из званичне музејске базе података.\n')
notes.add_run('• ').bold = True
notes.add_run('Препоручује се додатна курирска провера пре финалне поставке.')

# Save document
output_filename = 'Stalna_Postavka_Mineraloske_Zbirke.docx'
doc.save(output_filename)

print(f'\n{"="*80}')
print(f'Документ успешно креиран: {output_filename}')
print(f'{"="*80}')
print(f'\nСтатистика:')
print(f'  • Укупно примерака: {len(balkan_specimens) + len(world_specimens)}')
print(f'  • Са Балкана: {len(balkan_specimens)}')
print(f'  • Из света: {len(world_specimens)}')
print(f'\nДокумент садржи детаљне информације о сваком примерку:')
print(f'  - Назив минерала')
print(f'  - Инвентарни број (када је доступан)')
print(f'  - Локалитет')
print(f'  - Додатне информације (опис, напомене)')
print(f'\n{"="*80}\n')

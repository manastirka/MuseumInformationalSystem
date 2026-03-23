#!/usr/bin/env python3
"""
Generate Museum Organizational Chart
Creates a .docx document with the complete organizational structure
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_border(paragraph, border_type='bottom'):
    """Add border to paragraph"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')

    if border_type == 'bottom':
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '2E5090')
        pBdr.append(bottom)

    pPr.append(pBdr)

def create_org_chart():
    """Create organizational chart document"""

    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('ПРИРОДЊАЧКИ МУЗЕЈ У БЕОГРАДУ')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(46, 80, 144)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('ОРГАНИЗАЦИОНА СТРУКТУРА')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(46, 80, 144)

    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle2.add_run('Organizational Chart with Subject Matter Experts')
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # Spacing

    # Organizational structure
    org_structure = {
        'ДИРЕКТОР / DIRECTOR': [
            {'name': 'MSc Славко Спасић', 'position': 'виши кустос', 'expertise': 'Museum Director, Senior Curator', 'email': 'slavko.spasic@nhmbeo.rs'}
        ],

        'ОДСЕК ОПШТИХ И ПРАВНИХ ПОСЛОВА / DEPARTMENT OF GENERAL AND LEGAL AFFAIRS': [
            {'name': 'Ана Живановић', 'position': 'секретар', 'expertise': 'Secretary, Legal Affairs', 'email': 'ana.zivanovic@nhmbeo.rs'},
            {'name': 'Ана Ковачевић', 'position': 'технички секретар', 'expertise': 'Technical Secretary', 'email': 'ana.kovacevic@nhmbeo.rs'},
            {'name': 'Бора Милићевић', 'position': 'препаратор – ликовни техничар', 'expertise': 'Preparator - Art Technician', 'email': 'bora.m@nhmbeo.rs'},
            {'name': 'Предраг Илић', 'position': 'препаратор – техничар графичке припреме', 'expertise': 'Preparator - Graphic Preparation Technician', 'email': 'pedja@nhmbeo.rs'},
            {'name': 'Мср Оливера Аломеровић', 'position': 'дипл. библиотекар', 'expertise': 'Librarian, Information Science', 'email': 'biblioteka@nhmbeo.rs'},
            {'name': 'Данијела Мијаиловић', 'position': 'радник за одржавање хигијене', 'expertise': 'Hygiene Maintenance', 'email': '-'},
            {'name': 'Невенка Јанковић', 'position': 'радник за одржавање хигијене', 'expertise': 'Hygiene Maintenance', 'email': '-'}
        ],

        'ГРУПА ЗА ФИНАНСИЈСКО-РАЧУНОВОДСТВЕНЕ ПОСЛОВЕ / FINANCE AND ACCOUNTING GROUP': [
            {'name': 'MSc Душица Ивић', 'position': 'помоћник директора-руководилац одељења финансија', 'expertise': 'Assistant Director - Head of Finance Department', 'email': 'dusica.ivic@nhmbeo.rs'},
            {'name': 'Милена Радочај', 'position': 'финансијско-рачуноводствени референт', 'expertise': 'Finance and Accounting Clerk', 'email': 'milenar@nhmbeo.rs'},
            {'name': 'Милица Томић', 'position': 'финансијско-рачуноводствени сарадник', 'expertise': 'Finance and Accounting Associate', 'email': 'milica@nhmbeo.rs'}
        ],

        'ГРУПА ЗА ЕДУКАЦИЈУ, КОМУНИКАЦИЈУ И МАРКЕТИНГ / EDUCATION, COMMUNICATION AND MARKETING GROUP': [
            {'name': 'Дипл. инж. Драгана Вучићевић', 'position': 'виши кустос', 'expertise': 'Senior Curator, Education & Communication', 'email': 'draganav@nhmbeo.rs'},
            {'name': 'Мср Симка Вукојевић', 'position': 'музејски педагог', 'expertise': 'Museum Educator, Pedagogy', 'email': 'simka.vukojevic@nhmbeo.rs'}
        ],

        'ГРУПА ЗА ИЗЛОЖБЕНЕ ПОСЛОВЕ – ГАЛЕРИЈА / EXHIBITION WORK GROUP - GALLERY': [
            {'name': 'Милица Ракић', 'position': 'организаторка туристичке и услужне делатности', 'expertise': 'Tourism & Service Activities Organizer', 'email': 'milica.rakic@nhmbeo.rs'},
            {'name': 'Снежана Јовановић', 'position': 'водич', 'expertise': 'Museum Guide, Visitor Services', 'email': 'galerija@nhmbeo.rs'}
        ],

        'БИОЛОШКО ОДЕЉЕЊЕ / BIOLOGY DEPARTMENT': [
            {'name': 'Др Марјан Никетић', 'position': 'руководилац Биолошког одељења, музејски саветник ботаничар', 'expertise': 'Head of Biology Dept, Museum Advisor - Botanist', 'email': 'mniketic@nhmbeo.rs'},
            {'name': 'Дубравка Вучић', 'position': 'музејски саветник ихтиолог', 'expertise': 'Museum Advisor - Ichthyologist (Fish)', 'email': 'dubravka.vucic@nhmbeo.rs'},
            {'name': 'Милош Јовић', 'position': 'музејски саветник ентомолог', 'expertise': 'Museum Advisor - Entomologist (Insects)', 'email': 'milos.jovic@nhmbeo.rs'},
            {'name': 'Др Борис Иванчевић', 'position': 'музејски саветник миколог', 'expertise': 'Museum Advisor - Mycologist (Fungi)', 'email': 'boris@nhmbeo.rs'},
            {'name': 'Др Ана Пауновић', 'position': 'музејски саветник херпетолог', 'expertise': 'Museum Advisor - Herpetologist (Reptiles & Amphibians)', 'email': 'ana.paunovic@nhmbeo.rs'},
            {'name': 'Дипл.инж. Александар Стојановић', 'position': 'конзерватор ентомолог', 'expertise': 'Conservator - Entomologist', 'email': 'aleksandar@nhmbeo.rs'},
            {'name': 'Др Александра Савић', 'position': 'музејски саветник ботаничар / етноботаничар', 'expertise': 'Museum Advisor - Botanist / Ethnobotanist', 'email': 'aleksandra.savic@nhmbeo.rs'},
            {'name': 'Верица Стојановић', 'position': 'кустос приправник', 'expertise': 'Curator Trainee', 'email': 'verica.stojanovic@nhmbeo.rs'},
            {'name': 'Горана Петковски', 'position': 'конзерватор', 'expertise': 'Conservator', 'email': 'gorana.petkovski@nhmbeo.rs'},
            {'name': 'Др Марко Несторовић', 'position': 'музејски саветник ботаничар / херболог', 'expertise': 'Museum Advisor - Botanist / Herbologist', 'email': 'marko.nestorovic@nhmbeo.rs'},
            {'name': 'Зорана Марковић', 'position': 'кустос', 'expertise': 'Curator', 'email': 'zorana.markovic@nhmbeo.rs'},
            {'name': 'Мср Вук Попић', 'position': 'кустос приправник', 'expertise': 'Curator Trainee', 'email': 'vuk.popic@nhmbeo.rs'},
            {'name': 'Милош Мрваљевић', 'position': 'препаратор приправник', 'expertise': 'Preparator Trainee', 'email': 'milos.mrvaljevic@nhmbeo.rs'},
            {'name': 'Јован Кокотовић', 'position': 'кустос приправник', 'expertise': 'Curator Trainee', 'email': 'jovan.kokotovic@nhmbeo.rs'}
        ],

        'ГЕОЛОШКО ОДЕЉЕЊЕ / GEOLOGY DEPARTMENT': [
            {'name': 'Др Биљана Митровић', 'position': 'начелник Геолошког одељења, музејски саветник палеозоолог', 'expertise': 'Head of Geology Dept, Museum Advisor - Paleozoologist', 'email': 'biljana.mitrovic@nhmbeo.rs'},
            {'name': 'Др Зоран Марковић', 'position': 'музејски саветник палеозоолог', 'expertise': 'Museum Advisor - Paleozoologist', 'email': 'zoran.markovic@nhmbeo.rs'},
            {'name': 'Мср Сања Алабурић', 'position': 'музејски саветник палеозоолог', 'expertise': 'Museum Advisor - Paleozoologist', 'email': 'sanja.pavic@nhmbeo.rs'},
            {'name': 'Др Александра Маран Стевановић', 'position': 'музејски саветник палеозоолог', 'expertise': 'Museum Advisor - Paleozoologist', 'email': 'amaran@nhmbeo.rs'},
            {'name': 'Др Деса Ђорђевић-Милутиновић', 'position': 'музејски саветник, палеоботаничар', 'expertise': 'Museum Advisor - Paleobotanist', 'email': 'desadjm@nhmbeo.rs'},
            {'name': 'Др Александар Луковић', 'position': 'кустос минералог', 'expertise': 'Curator - Mineralogist, Database Systems', 'email': 'aca.lukovic@nhmbeo.rs'},
            {'name': 'Др Драгана Ђурић', 'position': 'музејски саветник палеозоолог', 'expertise': 'Museum Advisor - Paleozoologist', 'email': 'dragana.djuric@nhmbeo.rs'},
            {'name': 'Татјана Милић Бабић', 'position': 'виши кустос петролог', 'expertise': 'Senior Curator - Petrologist', 'email': 'tatjana.milicbabic@nhmbeo.rs'},
            {'name': 'Дипл. инж. Ранко Пејовић', 'position': 'кустос палеозоолог', 'expertise': 'Curator - Paleozoologist', 'email': 'pejovic.ranko@nhmbeo.rs'},
            {'name': 'Милош Миливојевић', 'position': 'виши препаратор за геолошке збирке', 'expertise': 'Senior Preparator for Geological Collections', 'email': 'milos.milivojevic@nhmbeo.rs'},
            {'name': 'Мср Бранко Радуловић', 'position': 'кустос за геолошке збирке', 'expertise': 'Curator for Geological Collections', 'email': 'branko.radulovic@nhmbeo.rs'},
            {'name': 'Мср Ненад Младеновић', 'position': 'конзерватор приправник', 'expertise': 'Conservator Trainee', 'email': 'nenad.mladenovic@nhmbeo.rs'}
        ]
    }

    # Add each department
    for dept_name, employees in org_structure.items():
        # Department header
        dept_para = doc.add_paragraph()
        dept_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = dept_para.add_run(dept_name)
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(46, 80, 144)
        add_border(dept_para, 'bottom')

        # Add spacing
        doc.add_paragraph()

        # Add employees
        for emp in employees:
            # Employee name
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(f"• {emp['name']}")
            name_run.font.size = Pt(11)
            name_run.font.bold = True
            name_run.font.color.rgb = RGBColor(0, 0, 0)

            # Position (Serbian)
            pos_para = doc.add_paragraph()
            pos_para.paragraph_format.left_indent = Inches(0.3)
            pos_run = pos_para.add_run(f"   Позиција: {emp['position']}")
            pos_run.font.size = Pt(10)
            pos_run.font.italic = True
            pos_run.font.color.rgb = RGBColor(80, 80, 80)

            # Expertise (English)
            exp_para = doc.add_paragraph()
            exp_para.paragraph_format.left_indent = Inches(0.3)
            exp_run = exp_para.add_run(f"   Expertise: {emp['expertise']}")
            exp_run.font.size = Pt(10)
            exp_run.font.color.rgb = RGBColor(60, 60, 60)

            # Email
            email_para = doc.add_paragraph()
            email_para.paragraph_format.left_indent = Inches(0.3)
            email_run = email_para.add_run(f"   Email: {emp['email']}")
            email_run.font.size = Pt(9)
            email_run.font.color.rgb = RGBColor(100, 100, 100)

            # Small spacing between employees
            doc.add_paragraph()

        # Larger spacing between departments
        doc.add_paragraph()

    # Footer with summary
    doc.add_page_break()
    summary_title = doc.add_paragraph()
    summary_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = summary_title.add_run('САЖЕТАК / SUMMARY')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(46, 80, 144)

    doc.add_paragraph()

    # Department statistics
    stats = [
        ('ДИРЕКТОР / DIRECTOR', 1),
        ('ОДСЕК ОПШТИХ И ПРАВНИХ ПОСЛОВА / GENERAL & LEGAL AFFAIRS', 7),
        ('ГРУПА ЗА ФИНАНСИЈСКО-РАЧУНОВОДСТВЕНЕ ПОСЛОВЕ / FINANCE & ACCOUNTING', 3),
        ('ГРУПА ЗА ЕДУКАЦИЈУ, КОМУНИКАЦИЈУ И МАРКЕТИНГ / EDUCATION & MARKETING', 2),
        ('ГРУПА ЗА ИЗЛОЖБЕНЕ ПОСЛОВЕ – ГАЛЕРИЈА / EXHIBITION - GALLERY', 2),
        ('БИОЛОШКО ОДЕЉЕЊЕ / BIOLOGY DEPARTMENT', 14),
        ('ГЕОЛОШКО ОДЕЉЕЊЕ / GEOLOGY DEPARTMENT', 12)
    ]

    for dept, count in stats:
        stat_para = doc.add_paragraph()
        stat_run = stat_para.add_run(f"• {dept}: ")
        stat_run.font.size = Pt(11)
        count_run = stat_para.add_run(f"{count} employees")
        count_run.font.size = Pt(11)
        count_run.font.bold = True
        count_run.font.color.rgb = RGBColor(46, 80, 144)

    doc.add_paragraph()

    # Total
    total_para = doc.add_paragraph()
    total_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    total_run = total_para.add_run('УКУПНО / TOTAL: 41 ЗАПОСЛЕНИХ / EMPLOYEES')
    total_run.font.size = Pt(12)
    total_run.font.bold = True
    total_run.font.color.rgb = RGBColor(46, 80, 144)

    doc.add_paragraph()
    doc.add_paragraph()

    # Key expertise areas
    expertise_title = doc.add_paragraph()
    expertise_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = expertise_title.add_run('КЉУЧНЕ ОБЛАСТИ ЕКСПЕРТИЗЕ / KEY EXPERTISE AREAS')
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(46, 80, 144)

    doc.add_paragraph()

    expertise_areas = [
        ('БОТАНИКА / BOTANY', 'Dr Marjan Niketić, Dr Aleksandra Savić, Dr Marko Nestorović'),
        ('ЕНТОМОЛОГИЈА / ENTOMOLOGY', 'Miloš Jović, Aleksandar Stojanović'),
        ('ИХТИОЛОГИЈА / ICHTHYOLOGY', 'Dubravka Vučić'),
        ('МИКОЛОГИЈА / MYCOLOGY', 'Dr Boris Ivančević'),
        ('ХЕРПЕТОЛОГИЈА / HERPETOLOGY', 'Dr Ana Paunović'),
        ('ПАЛЕОЗООЛОГИЈА / PALEOZOOLOGY', 'Dr Biljana Mitrović, Dr Zoran Marković, Dr Aleksandra Maran Stevanović, Dr Dragana Đurić, Ranko Pejović'),
        ('ПАЛЕОБОТАНИКА / PALEOBOTANY', 'Dr Desa Đorđević-Milutinović'),
        ('МИНЕРАЛОГИЈА / MINERALOGY', 'Dr Aleksandar Luković'),
        ('ПЕТРОЛОГИЈА / PETROLOGY', 'Tatjana Milić Babić'),
        ('МУЗЕЈСКА ЕДУКАЦИЈА / MUSEUM EDUCATION', 'Dragana Vučićević, Simka Vukojević'),
        ('БИБЛИОТЕКАРСТВО / LIBRARY SCIENCE', 'Olivera Alomerović'),
        ('КОНЗЕРВАЦИЈА / CONSERVATION', 'Gorana Petkovski, Nenad Mladenović')
    ]

    for area, experts in expertise_areas:
        area_para = doc.add_paragraph()
        area_run = area_para.add_run(f"• {area}")
        area_run.font.size = Pt(11)
        area_run.font.bold = True
        area_run.font.color.rgb = RGBColor(46, 80, 144)

        experts_para = doc.add_paragraph()
        experts_para.paragraph_format.left_indent = Inches(0.3)
        experts_run = experts_para.add_run(f"   {experts}")
        experts_run.font.size = Pt(10)
        experts_run.font.color.rgb = RGBColor(60, 60, 60)

    # Save document
    filename = 'Museum_Organizational_Chart.docx'
    doc.save(filename)
    print(f"\n✅ Organizational chart created successfully!")
    print(f"📄 File saved as: {filename}")
    print(f"📊 Total employees documented: 41")
    print(f"🏢 Total departments: 7")

    return filename

if __name__ == "__main__":
    create_org_chart()

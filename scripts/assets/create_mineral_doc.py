#!/usr/bin/env python3
"""
Create new mineral collection document using:
  - Locality descriptions from the original document
  - ALL images from 'Slike sa razmernikom' and 'spisak i fotografije druza minerala'
  - Table format with 3 columns, polished design
  - Images resized to 600px max dimension before embedding (small file size)
  - Keep each item on same page (cantSplit rows)
  - No duplicate inventory numbers
"""

import os
import re
import io
import tempfile
from lxml import etree
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage

BASE = "/home/aleksandarlukovic/Desktop/DokumentacijaNovaZgrada/Stalna postavka/Geologija/Predlozi eksponata zaposlenih u geoloskom odeljenju/Aleksandar"
SRC_DOC = os.path.join(BASE, "Aleksandar Luković-Zbirka minerala deo 1, rudnici– Stalna postavka_SA_SLIKAMA.docx")
OUT_DOC = os.path.join(BASE, "Aleksandar Luković-Zbirka minerala sa slikama.docx")

SCALE_DIR = os.path.join(BASE, "Slike sa razmernikom")
DRUZA_DIR = os.path.join(BASE, "spisak i fotografije druza minerala")

# Display size for images in cells
IMG_DISPLAY_CM = 4.6
# Max pixel dimension for embedded images (600px at 4.6cm ≈ 330dpi - sharp enough for print)
IMG_MAX_PX = 600
# JPEG quality for embedded images
IMG_QUALITY = 82
# Column width in twips (1/20 of a point). 3 cols in ~16.6cm usable = ~5.53cm each
COL_WIDTH_DXA = 3150
# Temp dir for resized images
TEMP_DIR = None

# Colors
CLR_DARK = RGBColor(0x1A, 0x2B, 0x3C)    # Dark navy for headers
CLR_ACCENT = RGBColor(0xB7, 0x2B, 0x2B)   # Deep red for inv numbers
CLR_MUTED = RGBColor(0x6B, 0x7B, 0x8D)    # Muted grey-blue for locality
CLR_SECTION = RGBColor(0x2C, 0x3E, 0x50)   # Section headers
CLR_BORDER = RGBColor(0xDD, 0xDD, 0xDD)    # Light border

LOCALITY_ORDER = [
    'trepca', 'kopaonik', 'mezice', 'busovaca', 'stolice',
    'bor', 'zlatibor', 'zajaca', 'rudnik', 'lece', 'avala',
    'majdanpek', 'alsar', 'presevo', 'italija',
]

FILENAME_LOCALITY_MAP = {
    'Suvo Rudište': 'kopaonik', 'Kopaonik': 'kopaonik',
    'Novo Brdo': 'kopaonik', 'Zimovnik': 'kopaonik',
    'Mežice': 'mezice', 'Mezice': 'mezice',
    'Busovač': 'busovaca', 'Busovac': 'busovaca',
    'Tisovac': 'busovaca', 'Kapak-Polom': 'busovaca',
    'Stolice': 'stolice', 'Stolica': 'stolice', 'Krupanj': 'stolice',
    'Bor, S': 'bor',
    'Zlatibor': 'zlatibor', 'Krnda': 'zlatibor',
    'Zajač': 'zajaca',
    'Rudnik': 'rudnik',
    'Lece': 'lece',
    'Avala': 'avala', 'Leštane': 'avala', 'Ljubičica': 'avala',
    'Majdanpek': 'majdanpek',
    'Alšar': 'alsar', 'Makedonij': 'alsar',
    'Preševo': 'presevo',
    'Traversella': 'italija', 'Piemonte': 'italija',
    'Trepča': 'trepca', 'Trepca': 'trepca', 'Stari Trg': 'trepca',
}


def guess_locality(filename):
    for pattern, key in FILENAME_LOCALITY_MAP.items():
        if pattern in filename:
            return key
    return None


def extract_name_from_filename(filename):
    name = os.path.splitext(filename)[0]
    name = re.sub(r'^\d+[\s,-]+\d*[\s,-]*', '', name)
    name = re.sub(r'^\d+\s+\d+,\s*', '', name)
    name = re.sub(r'[_]Trepča$', '', name)
    name = re.sub(r'^_', '', name)
    name = re.sub(r'\s+\d+\.?\s*$', '', name)

    loc_words = {'Suvo', 'Rudište', 'Kopaonik', 'Novo', 'Brdo', 'Zimovnik',
                 'Mežice', 'Mezice', 'Busovača', 'Busovaca', 'Tisovac',
                 'Kapak', 'Polom', 'Stolice', 'Stolica', 'Krupanj', 'Krupnja',
                 'Bor', 'Zlatibor', 'Krnda', 'Zajača', 'Rudnik', 'Lece',
                 'Avala', 'Leštane', 'Ljubičica', 'Majdanpek', 'Alšar',
                 'Makedonija', 'Preševo', 'Traversella', 'Piemonte',
                 'Trepča', 'Stari', 'Srbija', 'Bosna', 'BiH', 'Slovenija',
                 'Italija', 'Kosovo', 'Metohija', 'Potkop'}

    parts = [p.strip() for p in name.split(',')]
    mineral_parts = []
    for part in parts:
        words = part.split()
        if any(w in loc_words for w in words):
            break
        mineral_parts.append(part)

    result = ', '.join(mineral_parts).strip(' ,.')
    if re.match(r'^DSC\d+$', result):
        result = 'МИНЕРАЛ ' + result
    return result.upper() if result else filename.upper()


def extract_locality_from_filename(filename):
    name = os.path.splitext(filename)[0]
    name = re.sub(r'^\d+[\s,-]+\d*[\s,-]*', '', name)
    name = re.sub(r'^\d+\s+\d+,\s*', '', name)
    name = re.sub(r'^_', '', name)

    loc_words = {'Suvo', 'Rudište', 'Kopaonik', 'Novo', 'Brdo', 'Zimovnik',
                 'Mežice', 'Mezice', 'Busovača', 'Busovaca', 'Tisovac',
                 'Kapak', 'Polom', 'Stolice', 'Stolica', 'Krupanj', 'Krupnja',
                 'Bor', 'Zlatibor', 'Krnda', 'Zajača', 'Rudnik', 'Lece',
                 'Avala', 'Leštane', 'Ljubičica', 'Majdanpek', 'Alšar',
                 'Makedonija', 'Preševo', 'Traversella', 'Piemonte',
                 'Trepča', 'Stari', 'Srbija', 'Bosna', 'BiH', 'Slovenija',
                 'Italija', 'Potkop'}

    parts = [p.strip() for p in name.split(',')]
    loc_parts = []
    found_loc = False
    for part in parts:
        words = part.split()
        if any(w in loc_words for w in words):
            found_loc = True
        if found_loc:
            loc_parts.append(part)

    if loc_parts:
        return ', '.join(loc_parts).strip(' ,.')
    if '_Trepča' in filename or 'Trepča' in filename:
        return 'Stari Trg, Trepča, Srbija'
    return ''


def resize_image(img_path):
    """Resize image to max IMG_MAX_PX and save as JPEG in temp dir.
    Returns path to resized image and (width, height) in pixels."""
    with PILImage.open(img_path) as im:
        w, h = im.size

        # Calculate new size
        if max(w, h) > IMG_MAX_PX:
            if w >= h:
                new_w = IMG_MAX_PX
                new_h = int(h * IMG_MAX_PX / w)
            else:
                new_h = IMG_MAX_PX
                new_w = int(w * IMG_MAX_PX / h)
        else:
            new_w, new_h = w, h

        resized = im.resize((new_w, new_h), PILImage.LANCZOS)

        if resized.mode in ('RGBA', 'P'):
            resized = resized.convert('RGB')

        # Save to temp file
        tmp_path = os.path.join(TEMP_DIR, f'img_{id(img_path)}_{os.getpid()}.jpg')
        resized.save(tmp_path, 'JPEG', quality=IMG_QUALITY, optimize=True)
        return tmp_path, new_w, new_h


def collect_all_items():
    items = {}
    seen_inv = set()

    for root, dirs, files in os.walk(SCALE_DIR):
        for f in sorted(files):
            if not f.lower().endswith(('.jpg', '.jpeg', '.png')):
                continue
            fp = os.path.join(root, f)
            m = re.match(r'(\d+)', f)
            inv = m.group(1) if m else None
            locality_key = guess_locality(f)
            if not locality_key:
                locality_key = 'trepca'

            name = extract_name_from_filename(f)
            loc_str = extract_locality_from_filename(f)

            if inv:
                if inv in seen_inv:
                    continue
                seen_inv.add(inv)
                key = f'inv_{inv}'
            else:
                key = f'scale_{f}'

            if not loc_str:
                loc_str = 'Stari Trg, Trepča, Srbija'

            items[key] = {
                'inv': inv, 'name': name, 'locality': loc_str,
                'locality_key': locality_key, 'img_path': fp, 'source': 'scale',
            }

    druza_groups = {}
    for f in sorted(os.listdir(DRUZA_DIR)):
        if not f.lower().endswith(('.jpg', '.jpeg', '.png')):
            continue
        fp = os.path.join(DRUZA_DIR, f)
        m = re.match(r'(\d+)', f)
        if m:
            inv = m.group(1)
            if inv not in druza_groups:
                druza_groups[inv] = []
            druza_groups[inv].append((f, fp))
        else:
            lm = re.search(r'bez br\.\s*([a-g])', f, re.IGNORECASE)
            if lm:
                label = lm.group(1).lower()
                key = f'bez_br_{label}'
                if key not in druza_groups:
                    druza_groups[key] = []
                druza_groups[key].append((f, fp))

    for group_key, file_list in druza_groups.items():
        best = file_list[0]
        for f, fp in file_list:
            if re.search(r'\b1[.,]', f) or re.search(r'1\.a\.', f):
                best = (f, fp)
                break
        f, fp = best

        if group_key.startswith('bez_br_'):
            inv = None
            label = group_key.split('_')[-1].upper()
            name = f'ДРУЗА БЕЗ БРОЈА ({label})'
            item_key = group_key
        else:
            inv = group_key
            if inv in seen_inv:
                continue
            seen_inv.add(inv)
            name = extract_name_from_filename(f)
            item_key = f'inv_{inv}'

        items[item_key] = {
            'inv': inv, 'name': name,
            'locality': 'Stari Trg, Trepča, Srbija',
            'locality_key': 'trepca', 'img_path': fp, 'source': 'druza',
        }

    return items


def extract_locality_descriptions(doc):
    locality_keywords = {
        'ТРЕПЧА': 'trepca', 'КОПАОНИК': 'kopaonik', 'МЕЖИЦА': 'mezice',
        'БУСОВАЧА': 'busovaca', 'СТОЛИЦЕ': 'stolice', 'БОР': 'bor',
        'ЗАЈАЧА': 'zajaca', 'РУДНИК': 'rudnik', 'ЛЕЦЕ': 'lece',
        'АВАЛА': 'avala', 'МАЈДАНПЕК': 'majdanpek', 'АЛШАР': 'alsar',
    }

    sections = {}
    current_key = None
    collecting = False

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = p.style.name if p.style else ''

        is_header = False
        if style == 'List Paragraph':
            for cyrillic, key in locality_keywords.items():
                if cyrillic in text and 'Инв. број' not in text:
                    current_key = key
                    if key not in sections:
                        sections[key] = {'header': text, 'descriptions': []}
                    collecting = True
                    is_header = True
                    break

        if is_header:
            continue

        if collecting and current_key:
            if 'Примерци из збирке' in text or 'Укупно примерака' in text:
                collecting = False
                continue
            if style == 'List Paragraph':
                collecting = False
                continue
            sections[current_key]['descriptions'].append(text)

    intro = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if p.style and p.style.name == 'List Paragraph':
            break
        intro.append(text)

    lit = []
    found_lit = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if 'ЛИТЕРАТУРА И НАПОМЕНЕ' in text:
            found_lit = True
        if found_lit and text:
            lit.append(text)

    return sections, intro, lit


# ============================================================
# Document building helpers
# ============================================================

def set_paragraph_spacing(paragraph, before=0, after=0, line=None):
    """Set paragraph spacing in twips."""
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = p_pr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        p_pr.append(spacing)
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    if line is not None:
        spacing.set(qn('w:line'), str(line))
        spacing.set(qn('w:lineRule'), 'auto')


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement('w:cantSplit')
    tr_pr.append(cant_split)


def set_keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_pr.append(OxmlElement('w:keepNext'))


def set_keep_lines(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_pr.append(OxmlElement('w:keepLines'))


def set_cell_margins(cell, top=40, bottom=40, left=60, right=60):
    """Set cell margins in twips."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('start', left), ('end', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tc_mar.append(el)
    tc_pr.append(tc_mar)


def set_cell_border_bottom(cell, color='DDDDDD', size=4):
    """Add a subtle bottom border to a cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), color)
    borders.append(bottom)
    tc_pr.append(borders)


def add_image_to_paragraph(paragraph, img_path):
    """Add resized image to paragraph, maintaining aspect ratio."""
    tmp_path, px_w, px_h = resize_image(img_path)

    max_dim = Cm(IMG_DISPLAY_CM)
    if px_w >= px_h:
        width = max_dim
        height = int(max_dim * px_h / px_w)
    else:
        height = max_dim
        width = int(max_dim * px_w / px_h)

    run = paragraph.add_run()
    run.add_picture(tmp_path, width=width, height=height)

    # Clean up temp file
    try:
        os.remove(tmp_path)
    except OSError:
        pass


def fill_cell(cell, item):
    """Fill a table cell with image + name + M-number + locality."""
    for p in cell.paragraphs:
        p.clear()

    # Cell margins
    set_cell_margins(cell, top=60, bottom=80, left=40, right=40)

    # Image
    p_img = cell.paragraphs[0]
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p_img, before=0, after=40)
    set_keep_lines(p_img)
    img_path = item['img_path']
    if img_path and os.path.exists(img_path):
        add_image_to_paragraph(p_img, img_path)

    # Name (bold, dark navy)
    p_name = cell.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p_name, before=20, after=0, line=240)
    run = p_name.add_run(item['name'])
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = CLR_DARK

    # M-number (bold, red accent)
    if item.get('inv'):
        p_inv = cell.add_paragraph()
        p_inv.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p_inv, before=0, after=0, line=220)
        run = p_inv.add_run(f"M{item['inv']}")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = CLR_ACCENT

    # Locality (muted grey)
    if item.get('locality'):
        p_loc = cell.add_paragraph()
        p_loc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p_loc, before=0, after=0, line=200)
        run = p_loc.add_run(item['locality'])
        run.font.size = Pt(7)
        run.font.color.rgb = CLR_MUTED


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(width_dxa))
    tc_w.set(qn('w:type'), 'dxa')


def add_horizontal_line(doc, color='2C3E50', width=8):
    """Add a thin decorative horizontal line."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=60, after=60)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(width))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    return p


def add_section_header(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=120, after=60)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = CLR_SECTION
    set_keep_with_next(p)
    # Decorative line under header
    add_horizontal_line(doc, color='2C3E50', width=6)
    return p


def add_description(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, before=40, after=40, line=276)

    lines = text.split('\n')
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
        if i > 0:
            p.add_run('\n')
        bold_prefixes = ['Регион:', 'Тип минерализације:', 'Геолошки контекст:',
                         'Историјски значај:', 'Минералошке особености:']
        made_bold = False
        for prefix in bold_prefixes:
            if line.startswith(prefix):
                run_b = p.add_run(prefix + ' ')
                run_b.bold = True
                run_b.font.size = Pt(9.5)
                run_b.font.color.rgb = CLR_SECTION
                run_r = p.add_run(line[len(prefix):].strip())
                run_r.font.size = Pt(9.5)
                run_r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                made_bold = True
                break
        if not made_bold:
            run = p.add_run(line)
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    set_keep_with_next(p)
    return p


def create_items_table(doc, items_list):
    if not items_list:
        return

    num_rows = (len(items_list) + 2) // 3
    table = doc.add_table(rows=num_rows, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Table borders: subtle horizontal separators only
    tbl_pr = table._tbl.tblPr
    tbl_borders = OxmlElement('w:tblBorders')
    for bname in ['top', 'left', 'bottom', 'right', 'insideV']:
        border = OxmlElement(f'w:{bname}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tbl_borders.append(border)
    # Subtle horizontal lines between rows
    insideH = OxmlElement('w:insideH')
    insideH.set(qn('w:val'), 'single')
    insideH.set(qn('w:sz'), '4')
    insideH.set(qn('w:space'), '0')
    insideH.set(qn('w:color'), 'E8E8E8')
    tbl_borders.append(insideH)
    tbl_pr.append(tbl_borders)

    # Table cell spacing for gap between columns
    tbl_cell_spacing = OxmlElement('w:tblCellSpacing')
    tbl_cell_spacing.set(qn('w:w'), '30')
    tbl_cell_spacing.set(qn('w:type'), 'dxa')
    tbl_pr.append(tbl_cell_spacing)

    for ri in range(num_rows):
        row = table.rows[ri]
        set_row_cant_split(row)

        for ci in range(3):
            cell = row.cells[ci]
            set_cell_width(cell, COL_WIDTH_DXA)

            # Vertical align top
            tc_pr = cell._tc.get_or_add_tcPr()
            v_align = OxmlElement('w:vAlign')
            v_align.set(qn('w:val'), 'top')
            tc_pr.append(v_align)

            idx = ri * 3 + ci
            if idx < len(items_list):
                fill_cell(cell, items_list[idx])

    return table


def main():
    global TEMP_DIR
    TEMP_DIR = tempfile.mkdtemp(prefix='mineral_doc_')

    print("Reading original document for locality descriptions...")
    orig_doc = Document(SRC_DOC)
    sections, intro_paras, lit_paras = extract_locality_descriptions(orig_doc)
    print(f"  Extracted {len(sections)} locality sections")

    print("Collecting all images...")
    all_items = collect_all_items()
    print(f"  Total unique items: {len(all_items)}")

    items_by_locality = {}
    for item in all_items.values():
        key = item['locality_key']
        if key not in items_by_locality:
            items_by_locality[key] = []
        items_by_locality[key].append(item)

    for key in items_by_locality:
        items_by_locality[key].sort(
            key=lambda x: (0 if x['inv'] else 1, int(x['inv']) if x['inv'] else 99999, x['name'])
        )

    print("  Items per locality:")
    for key in LOCALITY_ORDER:
        if key in items_by_locality:
            print(f"    {key}: {len(items_by_locality[key])} items")

    # Create document
    print("Creating document...")
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # === TITLE PAGE ===
    # Add some spacing before title
    for _ in range(3):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=120)
    run = p.add_run('МИНЕРАЛОШКА БОГАТСТВА ЈУГОСЛАВИЈЕ')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = CLR_SECTION

    add_horizontal_line(doc, color='B72B2B', width=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=80, after=200)
    run = p.add_run('Стална поставка — Чувена рудишта и локалитети')
    run.font.size = Pt(13)
    run.font.color.rgb = CLR_MUTED

    # Intro paragraphs
    for text in intro_paras:
        if text in ('УВОД', 'ДИМЕНЗИЈЕ', 'САДРЖАЈ'):
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=160, after=40)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = CLR_SECTION
        elif text.startswith('ТРЕПЧА') and 'КОПАОНИК' in text:
            pass  # Skip old summary
        elif len(text) > 20:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(p, before=20, after=40, line=276)
            run = p.add_run(text)
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Contents summary
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=200, after=60)
    run = p.add_run('САДРЖАЈ ПОСТАВКЕ')
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = CLR_SECTION

    add_horizontal_line(doc, color='2C3E50', width=4)

    names = {
        'trepca': 'ТРЕПЧА — Стари Трг (Косово и Метохија, Србија)',
        'kopaonik': 'КОПАОНИК (Централна Србија)',
        'mezice': 'МЕЖИЦА (Словенија)',
        'busovaca': 'БУСОВАЧА — Тисовац (Босна и Херцеговина)',
        'stolice': 'СТОЛИЦЕ (Западна Србија)',
        'bor': 'БОР (Источна Србија)',
        'zlatibor': 'ЗЛАТИБОР (Западна Србија)',
        'zajaca': 'ЗАЈАЧА (Западна Србија)',
        'rudnik': 'РУДНИК (Централна Србија)',
        'lece': 'ЛЕЦЕ (Јужна Србија)',
        'avala': 'АВАЛА (Београд, Србија)',
        'majdanpek': 'МАЈДАНПЕК (Источна Србија)',
        'alsar': 'АЛШАР (Северна Македонија)',
        'presevo': 'ПРЕШЕВО (Јужна Србија)',
        'italija': 'ТРАВEРСEЛА (Пијемонт, Италија)',
    }

    for key in LOCALITY_ORDER:
        if key not in items_by_locality:
            continue
        count = len(items_by_locality[key])
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=20, after=20)
        run = p.add_run(f"  {names.get(key, key)}")
        run.font.size = Pt(10)
        run.font.color.rgb = CLR_SECTION
        run_count = p.add_run(f"  —  {count} примерака")
        run_count.font.size = Pt(10)
        run_count.font.color.rgb = CLR_MUTED

    # === LOCALITY SECTIONS ===
    total_items = 0
    for loc_idx, key in enumerate(LOCALITY_ORDER):
        if key not in items_by_locality:
            continue

        items_list = items_by_locality[key]
        total_items += len(items_list)

        doc.add_page_break()

        section_data = sections.get(key)
        if section_data:
            add_section_header(doc, section_data['header'])
            for desc_text in section_data['descriptions']:
                add_description(doc, desc_text)
        else:
            header_names = {
                'zlatibor': 'ЗЛАТИБОР',
                'presevo': 'ПРЕШЕВО',
                'italija': 'ТРАВEРСEЛА, Италија',
            }
            add_section_header(doc, header_names.get(key, key.upper()))

        # Specimen count
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=60, after=80)
        run = p.add_run(f'Примерака у поставци: {len(items_list)}')
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = CLR_MUTED

        # Items table
        print(f"  {key}: {len(items_list)} items...")
        create_items_table(doc, items_list)

    # === LITERATURE ===
    doc.add_page_break()
    for text in lit_paras:
        if text == 'ЛИТЕРАТУРА И НАПОМЕНЕ':
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=120, after=60)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = CLR_SECTION
            add_horizontal_line(doc, color='2C3E50', width=4)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(p, before=20, after=20, line=276)
            run = p.add_run(text)
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    print(f"\nTotal items: {total_items}")
    print(f"Saving: {OUT_DOC}")
    doc.save(OUT_DOC)

    # Clean up temp dir
    import shutil
    try:
        shutil.rmtree(TEMP_DIR)
    except OSError:
        pass

    fsize = os.path.getsize(OUT_DOC)
    print(f"File size: {fsize / 1024 / 1024:.1f} MB")
    print("Done!")


if __name__ == '__main__':
    main()

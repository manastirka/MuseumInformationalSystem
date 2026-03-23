#!/usr/bin/env python3
"""
Extract exhibition data from Word document and prepare for database import
"""
import docx
import json
import re
from datetime import datetime

def parse_date(date_str):
    """Parse various date formats from Serbian text"""
    if not date_str:
        return None
    
    # Clean up the string
    date_str = date_str.strip()
    
    # Common patterns
    # Example: "15. мај - 15. јун 2015"
    # Example: "01.06.2015 - 30.06.2015"
    # Example: "март 2015"
    
    return date_str  # Return as-is for now, will format later

def extract_exhibitions_from_docx(docx_path):
    """Extract exhibition data from Word document"""
    doc = docx.Document(docx_path)
    
    exhibitions = []
    current_exhibition = {}
    
    print(f"📄 Reading document: {docx_path}")
    print(f"   Total paragraphs: {len(doc.paragraphs)}")
    print(f"   Total tables: {len(doc.tables)}")
    print()
    
    # Extract text from all paragraphs
    full_text = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)
            print(f"   {text[:100]}...")
    
    print(f"\n✅ Extracted {len(full_text)} paragraphs with text")
    
    # Also check tables
    for table_idx, table in enumerate(doc.tables):
        print(f"\n📊 Table {table_idx + 1}:")
        for row_idx, row in enumerate(table.rows):
            row_text = [cell.text.strip() for cell in row.cells]
            print(f"   Row {row_idx + 1}: {row_text}")
    
    return full_text, exhibitions

def main():
    """Main execution"""
    docx_path = "/home/aleksandarlukovic/MuseumInfoSystem/Prirodnjački muzej_ Izložbe zadnjih 15 godina.docx"
    
    print("🏛️ Extracting Exhibition Data from Word Document")
    print("=" * 60)
    
    text_lines, exhibitions = extract_exhibitions_from_docx(docx_path)
    
    # Save raw text to file for review
    output_txt = "/home/aleksandarlukovic/MuseumInfoSystem/exhibitions_raw_text.txt"
    with open(output_txt, 'w', encoding='utf-8') as f:
        f.write('\n'.join(text_lines))
    
    print(f"\n📝 Raw text saved to: {output_txt}")
    print(f"   Total lines: {len(text_lines)}")

if __name__ == '__main__':
    main()

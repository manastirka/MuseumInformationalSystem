#!/usr/bin/env python3
"""
Import exhibition data from Word document into exhibitions.json database
"""
import docx
import json
import os
from datetime import datetime

def parse_exhibitions_from_table(table):
    """Parse exhibition data from Word table"""
    exhibitions = []
    
    # Skip header row
    for row_idx in range(1, len(table.rows)):
        row = table.rows[row_idx]
        cells = [cell.text.strip() for cell in row.cells]
        
        if len(cells) < 5:
            continue
        
        year = cells[0]
        title = cells[1]
        authors = cells[2]
        location = cells[3]
        themes = cells[4]
        
        # Skip if no title
        if not title or title == 'N/A':
            continue
        
        # Determine exhibition type
        if 'putujuća' in location.lower() or 'putujuća' in title.lower():
            ex_type = 'Путујућа изложба'
        else:
            ex_type = 'Привремена изложба'
        
        # Parse authors/curators
        curator = ''
        co_curator = ''
        if authors and authors != 'N/A':
            author_list = [a.strip() for a in authors.split(',')]
            if len(author_list) > 0:
                curator = author_list[0]
            if len(author_list) > 1:
                co_curator = ', '.join(author_list[1:])
        
        # Estimate dates based on year
        if year and year.isdigit():
            # Most exhibitions in the document run for several months
            start_date = f"{year}-01-01"  # Placeholder
            end_date = f"{year}-12-31"    # Placeholder
        else:
            start_date = None
            end_date = None
        
        exhibition = {
            'title': title,
            'title_en': '',  # To be filled manually
            'type': ex_type,
            'status': 'Завршена',
            'start_date': start_date,
            'end_date': end_date,
            'location': location,
            'curator': curator,
            'co_curator': co_curator,
            'specimens_count': 0,  # Unknown from document
            'species_count': 0,    # Unknown from document
            'visitor_count': 0,    # Unknown from document
            'description': f"Изложба са темама: {themes}",
            'description_en': '',
            'target_audience': 'Општа публика',
            'educational_programs': True,
            'guided_tours': True,
            'catalog_available': False,
            'keywords': themes.lower()
        }
        
        exhibitions.append(exhibition)
    
    return exhibitions

def load_existing_exhibitions(json_path):
    """Load existing exhibitions from JSON file"""
    if os.path.exists(json_path):
        with open(json_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    return []

def merge_exhibitions(existing, new_exhibitions):
    """Merge new exhibitions with existing, avoiding duplicates"""
    # Get existing titles
    existing_titles = {ex['title'].lower() for ex in existing}
    
    # Find next available ID
    max_id = max([ex['id'] for ex in existing]) if existing else 0
    
    merged = existing.copy()
    added_count = 0
    
    for ex in new_exhibitions:
        # Check if exhibition already exists
        if ex['title'].lower() not in existing_titles:
            max_id += 1
            ex['id'] = max_id
            merged.append(ex)
            added_count += 1
            print(f"  ➕ Added: {ex['title']}")
        else:
            print(f"  ⏭️  Skipped (duplicate): {ex['title']}")
    
    return merged, added_count

def main():
    """Main execution"""
    docx_path = "/home/aleksandarlukovic/MuseumInfoSystem/Prirodnjački muzej_ Izložbe zadnjih 15 godina.docx"
    json_path = "/home/aleksandarlukovic/MuseumInfoSystem/data/exhibitions.json"
    backup_path = "/home/aleksandarlukovic/MuseumInfoSystem/data/exhibitions_backup.json"
    
    print("🏛️ Importing Exhibition Data from Word Document")
    print("=" * 70)
    
    # Read Word document
    print(f"\n📄 Reading: {os.path.basename(docx_path)}")
    doc = docx.Document(docx_path)
    
    # Find and parse table
    if len(doc.tables) == 0:
        print("❌ No tables found in document")
        return
    
    print(f"📊 Found {len(doc.tables)} table(s)")
    
    # Parse exhibitions from first table
    new_exhibitions = parse_exhibitions_from_table(doc.tables[0])
    print(f"✅ Extracted {len(new_exhibitions)} exhibitions from table")
    
    # Load existing exhibitions
    print(f"\n📖 Loading existing database: {os.path.basename(json_path)}")
    existing_exhibitions = load_existing_exhibitions(json_path)
    print(f"   Current exhibitions: {len(existing_exhibitions)}")
    
    # Backup existing database
    if existing_exhibitions:
        with open(backup_path, 'w', encoding='utf-8') as f:
            json.dump(existing_exhibitions, f, ensure_ascii=False, indent=2)
        print(f"💾 Backup saved: {os.path.basename(backup_path)}")
    
    # Merge exhibitions
    print(f"\n🔄 Merging exhibitions...")
    merged_exhibitions, added_count = merge_exhibitions(existing_exhibitions, new_exhibitions)
    
    # Sort by ID
    merged_exhibitions.sort(key=lambda x: x['id'])
    
    # Save updated database
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(merged_exhibitions, f, ensure_ascii=False, indent=2)
    
    print(f"\n✅ Database updated successfully!")
    print(f"   Total exhibitions: {len(merged_exhibitions)}")
    print(f"   New additions: {added_count}")
    print(f"   Saved to: {json_path}")
    
    # Show summary of new exhibitions
    if added_count > 0:
        print(f"\n📋 New exhibitions added:")
        for ex in merged_exhibitions[-added_count:]:
            print(f"   {ex['id']}. {ex['title']} ({ex['start_date'][:4] if ex['start_date'] else 'N/A'})")

if __name__ == '__main__':
    main()

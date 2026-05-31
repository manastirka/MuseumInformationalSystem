"""Import Sanja's Paleogene/Neogene large mammal spreadsheet into JSON."""

from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path

import xlrd

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

SOURCE_PATH = Path("Sanja/sanja BAZA KRUPNI SISAri paleogen neogenNOVA.xls")
OUTPUT_PATH = Path("Sanja/sanja_paleogene_neogene_mammals.json")
LOCALITY_CODES_PATH = Path("Projekti/šifre lokaliteta.docx")

COLUMN_MAP = {
    0: "collection_number",
    1: "collector_number",
    2: "inventory_sequence",
    3: "catalog_number",
    4: "specimen_name",
    5: "taxon_name",
    6: "level",
    7: "acquisition_method",
    8: "quantity",
    9: "acquisition_date",
    10: "database_entry_date",
    11: "entered_by",
    12: "donor_collector",
    13: "identified_by",
    14: "comment",
    15: "description",
    16: "location_found",
    17: "utm_x",
    18: "utm_y",
    19: "utm_z",
    20: "country",
    21: "locality_text",
    22: "exposure",
    23: "specimen_part_2",
    24: "specimen_part_3",
    25: "specimen_part_4",
    26: "specimen_part_5",
    27: "storage_location",
    28: "material",
    29: "uncertain_determination",
    30: "length_mm",
    31: "width_mm",
    32: "biostratigraphic_zone",
    33: "bibliography_flag",
}

DATE_COLUMNS = {"acquisition_date", "database_entry_date"}


def normalize_cell(book: xlrd.book.Book, sheet: xlrd.sheet.Sheet, row: int, col: int):
    cell = sheet.cell(row, col)
    value = cell.value
    key = COLUMN_MAP[col]

    if value == "":
        return ""

    if key in DATE_COLUMNS and cell.ctype == xlrd.XL_CELL_DATE:
        dt = xlrd.xldate_as_datetime(value, book.datemode)
        if dt.time() == datetime.min.time():
            return dt.strftime("%Y-%m-%d")
        return dt.strftime("%Y-%m-%d %H:%M")

    if cell.ctype == xlrd.XL_CELL_NUMBER:
        if float(value).is_integer():
            return int(value)
        return round(float(value), 6)

    return str(value).strip()


def load_locality_codes() -> dict[str, str]:
    """Load locality code mapping from docx. Returns {locality_name: code}."""
    if DocxDocument is None or not LOCALITY_CODES_PATH.exists():
        return {}
    doc = DocxDocument(str(LOCALITY_CODES_PATH))
    code_map_raw = {}
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 2 and cells[0]:
                code_map_raw[cells[1]] = cells[0]

    # Build flexible mapping: base names without parenthetical, hyphen parts
    flexible = {}
    for name, code in code_map_raw.items():
        flexible[name] = code
        base = re.sub(r'\s*\(.*?\)\s*', '', name).strip()
        if base != name:
            flexible[base] = code
        parts = name.split('-')
        if len(parts) > 1:
            for p in parts:
                p = p.strip()
                if len(p) > 3:
                    flexible[p] = code
    # Also map parenthetical content -> code (e.g. "Zaječar" from "Trnavac (Zaječar)")
    for name, code in code_map_raw.items():
        m = re.search(r'\(([^)]+)\)', name)
        if m:
            flexible[m.group(1).strip()] = code
    return flexible


def build_records(book: xlrd.book.Book) -> list[dict]:
    locality_codes = load_locality_codes()
    sheet = book.sheet_by_index(0)
    records = []
    for row in range(1, sheet.nrows):
        record = {
            "id": row,
            "source_row": row + 1,
            "source_file": SOURCE_PATH.name,
            "collection_group": "Крупни сисари - палеоген/неоген",
        }
        for col in range(sheet.ncols):
            if col in COLUMN_MAP:
                record[COLUMN_MAP[col]] = normalize_cell(book, sheet, row, col)

        if not any(record.get(key) for key in ("catalog_number", "specimen_name", "location_found")):
            continue

        loc = record.get("location_found", "")
        record["locality_code"] = locality_codes.get(loc, "")

        records.append(record)
    return records


def build_payload(records: list[dict]) -> dict:
    localities = sorted({r["location_found"] for r in records if r.get("location_found")})
    taxa = sorted({r["specimen_name"] for r in records if r.get("specimen_name")})
    identified_by = sorted({r["identified_by"] for r in records if r.get("identified_by")})

    # Build locality code reference from docx
    locality_code_ref = {}
    if DocxDocument is not None and LOCALITY_CODES_PATH.exists():
        doc = DocxDocument(str(LOCALITY_CODES_PATH))
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 2 and cells[0]:
                    locality_code_ref[cells[0]] = cells[1]

    return {
        "metadata": {
            "name": "Крупни сисари палеогена и неогена",
            "source_file": str(SOURCE_PATH),
            "imported_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "locality_codes_source": str(LOCALITY_CODES_PATH),
            "locality_codes": locality_code_ref,
        },
        "specimens": records,
        "statistics": {
            "total_specimens": len(records),
            "total_taxa": len(taxa),
            "total_localities": len(localities),
            "identified_by_count": len(identified_by),
        },
    }


def main() -> None:
    book = xlrd.open_workbook(str(SOURCE_PATH))
    records = build_records(book)
    payload = build_payload(records)
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_PATH.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(f"Imported {len(records)} records to {OUTPUT_PATH}")


if __name__ == "__main__":
    main()

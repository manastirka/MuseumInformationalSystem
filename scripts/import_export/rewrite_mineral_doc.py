#!/usr/bin/env python3
"""
Rewrite 'Aleksandar Luković-Zbirka minerala sa slikama' document.
Uses images from:
  - Slike sa razmernikom (scale bar images)
  - spisak i fotografije druza minerala (druza photos)

Rules:
  - Remove items without images
  - No duplicate inventory numbers
  - Same tabular format (3 columns, centered, image + name + M-number + locality)
  - Druza items not in original doc are added to Trepča section
"""

import os
import re
import copy
import subprocess
from lxml import etree
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from PIL import Image as PILImage

BASE = "/home/aleksandarlukovic/Desktop/DokumentacijaNovaZgrada/Stalna postavka/Geologija/Predlozi eksponata zaposlenih u geoloskom odeljenju/Aleksandar"
SRC_DOC = os.path.join(BASE, "Aleksandar Luković-Zbirka minerala deo 1, rudnici– Stalna postavka_SA_SLIKAMA.docx")
OUT_DOC = os.path.join(BASE, "Aleksandar Luković-Zbirka minerala deo 1, rudnici– Stalna postavka_SA_SLIKAMA.docx")

SCALE_DIR = os.path.join(BASE, "Slike sa razmernikom")
DRUZA_DIR = os.path.join(BASE, "spisak i fotografije druza minerala")

IMG_SIZE_CM = 4.8
COL_WIDTH_DXA = 3005


def collect_images():
    """Collect all available images mapped by inventory number."""
    images = {}

    # Scale bar images (preferred)
    for root, dirs, files in os.walk(SCALE_DIR):
        for f in sorted(files):
            if not f.lower().endswith(('.jpg', '.jpeg', '.png')):
                continue
            m = re.match(r'(\d+)', f)
            if m:
                inv = m.group(1)
                if inv not in images:
                    images[inv] = os.path.join(root, f)

    # Druza images (use first/main view)
    for f in sorted(os.listdir(DRUZA_DIR)):
        if not f.lower().endswith(('.jpg', '.jpeg', '.png')):
            continue
        m = re.match(r'(\d+)', f)
        if not m:
            continue
        inv = m.group(1)
        if inv not in images:
            # Prefer image "1."
            if re.search(r'\b1[.,]', f):
                images[inv] = os.path.join(DRUZA_DIR, f)
            elif inv not in images:
                images[inv] = os.path.join(DRUZA_DIR, f)

    return images


def collect_druza_info():
    """Parse the druza RTF for names."""
    rtf_path = os.path.join(DRUZA_DIR, 'spisak druza po dimenzijama.rtf')
    txt_path = '/tmp/spisak_druza.txt'
    subprocess.run(['libreoffice', '--headless', '--convert-to', 'txt:Text',
                    rtf_path, '--outdir', '/tmp/'],
                   capture_output=True, timeout=30)
    converted = '/tmp/spisak druza po dimenzijama.txt'
    if os.path.exists(converted):
        os.rename(converted, txt_path)

    with open(txt_path, 'r', encoding='utf-8-sig') as f:
        text = f.read()

    entries = {}
    pattern = r'M(\d+)\s+(.+?)(?:\s*\([^)]*\))?\s*\n\s*D\s+'
    for m in re.finditer(pattern, text):
        inv = m.group(1)
        name = m.group(2).strip()
        entries[inv] = name
    return entries


def extract_existing_items(doc):
    """Extract all items from existing document tables."""
    items_by_table = {}
    for ti, table in enumerate(doc.tables):
        items = []
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if not text:
                    continue

                lines = [l.strip() for l in text.split('\n') if l.strip()]
                lines = [l for l in lines if l != 'Нема фотографије']

                m_num = None
                m_idx = None
                for i, l in enumerate(lines):
                    mm = re.match(r'^M(\d+)$', l)
                    if mm:
                        m_num = mm.group(1)
                        m_idx = i
                        break

                name = lines[0] if lines else ''
                name = re.sub(r'^\d+\.\s*', '', name)

                locality = ''
                if m_idx is not None and m_idx + 1 < len(lines):
                    locality = lines[m_idx + 1]
                elif len(lines) > 1:
                    locality = lines[-1]

                items.append({
                    'inv': m_num,
                    'name': name,
                    'locality': locality,
                })
        items_by_table[ti] = items
    return items_by_table


def add_image_to_paragraph(paragraph, img_path):
    """Add an image to a paragraph, maintaining aspect ratio within IMG_SIZE_CM."""
    run = paragraph.add_run()
    try:
        with PILImage.open(img_path) as im:
            w, h = im.size
    except Exception:
        w, h = 1, 1

    max_dim = Cm(IMG_SIZE_CM)
    if w >= h:
        width = max_dim
        height = int(max_dim * h / w)
    else:
        height = max_dim
        width = int(max_dim * w / h)

    run.add_picture(img_path, width=width, height=height)


def fill_cell(cell, name, inv_num, locality, img_path):
    """Fill a table cell with image + name + M-number + locality."""
    # Clear default paragraph
    for p in cell.paragraphs:
        p.clear()

    # Image
    if img_path and os.path.exists(img_path):
        p_img = cell.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_image_to_paragraph(p_img, img_path)

    # Name (bold, dark blue)
    p_name = cell.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run(name)
    run_name.bold = True
    run_name.font.size = Pt(8)
    run_name.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    # M-number (bold, red)
    if inv_num:
        p_inv = cell.add_paragraph()
        p_inv.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_inv = p_inv.add_run(f"M{inv_num}")
        run_inv.bold = True
        run_inv.font.size = Pt(9)
        run_inv.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    # Locality (grey)
    if locality:
        p_loc = cell.add_paragraph()
        p_loc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_loc = p_loc.add_run(locality)
        run_loc.font.size = Pt(7)
        run_loc.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)


def set_cell_width(cell, width_dxa):
    """Set cell width in DXA."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = etree.SubElement(tc_pr, qn('w:tcW'))
    tc_w.set(qn('w:w'), str(width_dxa))
    tc_w.set(qn('w:type'), 'dxa')


def main():
    print("Reading existing document...")
    orig_doc = Document(SRC_DOC)

    print("Collecting images...")
    images = collect_images()
    print(f"  {len(images)} specimens with images")

    druza_info = collect_druza_info()
    print(f"  {len(druza_info)} druza entries")

    existing_items = extract_existing_items(orig_doc)

    # Find existing inventory numbers
    existing_invs = set()
    for items in existing_items.values():
        for item in items:
            if item['inv']:
                existing_invs.add(item['inv'])

    # New druza items for Trepča section
    new_druza_items = []
    for inv in sorted(images.keys(), key=lambda x: int(x)):
        if inv not in existing_invs and inv in druza_info:
            name = druza_info[inv].upper()
            new_druza_items.append({
                'inv': inv,
                'name': name,
                'locality': 'Stari Trg, Trepča, Srbija',
            })
    print(f"  {len(new_druza_items)} new druza items")

    # Build output document - start fresh but copy from original
    print("Building document...")
    out_doc = Document(SRC_DOC)
    out_body = out_doc.element.body

    # Collect paragraphs and table positions from original
    orig_body = orig_doc.element.body
    elements_order = []  # ('para', elem) or ('table', table_idx)
    tbl_counter = 0
    for elem in orig_body:
        tag = elem.tag.split('}')[-1]
        if tag == 'tbl':
            elements_order.append(('table', tbl_counter))
            tbl_counter += 1
        elif tag == 'p':
            elements_order.append(('para', None))
        elif tag == 'sectPr':
            elements_order.append(('sectPr', None))

    # In the output doc, find and process each table
    out_tables = out_body.findall(qn('w:tbl'))
    global_seen = set()
    tables_to_remove = []

    for ti, tbl_elem in enumerate(out_tables):
        items_for_table = list(existing_items.get(ti, []))

        # Add druza items to table 0 (Trepča)
        if ti == 0:
            items_for_table.extend(new_druza_items)

        # Filter: only items with images, no duplicates
        valid_items = []
        for item in items_for_table:
            inv = item.get('inv')
            if not inv or inv in global_seen or inv not in images:
                continue
            valid_items.append(item)
            global_seen.add(inv)

        if not valid_items:
            tables_to_remove.append(tbl_elem)
            continue

        # Clear existing rows
        for tr in list(tbl_elem.findall(qn('w:tr'))):
            tbl_elem.remove(tr)

        # Build new rows (3 items per row)
        num_rows = (len(valid_items) + 2) // 3

        for ri in range(num_rows):
            tr_elem = etree.SubElement(tbl_elem, qn('w:tr'))

            for ci in range(3):
                idx = ri * 3 + ci
                tc_elem = etree.SubElement(tr_elem, qn('w:tc'))

                # Cell properties
                tc_pr = etree.SubElement(tc_elem, qn('w:tcPr'))
                tc_w = etree.SubElement(tc_pr, qn('w:tcW'))
                tc_w.set(qn('w:w'), str(COL_WIDTH_DXA))
                tc_w.set(qn('w:type'), 'dxa')

                # Empty paragraph (required)
                p_elem = etree.SubElement(tc_elem, qn('w:p'))
                p_pr = etree.SubElement(p_elem, qn('w:pPr'))
                jc = etree.SubElement(p_pr, qn('w:jc'))
                jc.set(qn('w:val'), 'center')

        print(f"  Table {ti}: {len(valid_items)} items")

    # Remove empty tables
    for tbl_elem in tables_to_remove:
        out_body.remove(tbl_elem)
        print(f"  Removed empty table")

    # Save structural document
    tmp_path = os.path.join(BASE, '_tmp_structural.docx')
    out_doc.save(tmp_path)

    # Reopen to fill cells with images via python-docx API
    print("Adding images to cells...")
    final_doc = Document(tmp_path)

    # Rebuild the mapping of tables to items
    global_seen2 = set()
    for ti_orig in range(len(out_tables)):
        items_for_table = list(existing_items.get(ti_orig, []))
        if ti_orig == 0:
            items_for_table.extend(new_druza_items)

        valid_items = []
        for item in items_for_table:
            inv = item.get('inv')
            if not inv or inv in global_seen2 or inv not in images:
                continue
            valid_items.append(item)
            global_seen2.add(inv)

        if not valid_items:
            continue

        # Find corresponding table in final doc
        # Tables may have shifted indices due to removals
        # We track by counting non-empty tables
        pass

    # Simpler: iterate final doc tables and fill them
    # We need to know which items go in which final table
    # Rebuild the ordered list of valid item groups
    global_seen3 = set()
    table_items_list = []  # list of valid_items per output table

    for ti_orig in range(len(out_tables)):
        items_for_table = list(existing_items.get(ti_orig, []))
        if ti_orig == 0:
            items_for_table.extend(new_druza_items)

        valid_items = []
        for item in items_for_table:
            inv = item.get('inv')
            if not inv or inv in global_seen3 or inv not in images:
                continue
            valid_items.append(item)
            global_seen3.add(inv)

        if valid_items:
            table_items_list.append(valid_items)

    # Now fill each table
    for t_idx, table in enumerate(final_doc.tables):
        if t_idx >= len(table_items_list):
            break
        valid_items = table_items_list[t_idx]

        for idx, item in enumerate(valid_items):
            ri = idx // 3
            ci = idx % 3
            if ri >= len(table.rows):
                break
            cell = table.rows[ri].cells[ci]
            inv = item['inv']
            img_path = images.get(inv)

            fill_cell(cell, item['name'], inv, item['locality'], img_path)

    os.remove(tmp_path)

    print(f"\nTotal items: {len(global_seen3)}")
    print(f"Saving: {OUT_DOC}")
    final_doc.save(OUT_DOC)
    print("Done!")


if __name__ == '__main__':
    main()

"""Тестови за стварни апп-извоз распоред (заглавље унутар табеле, транспоновано)
и за КОНКРЕТНЕ поруке одбијања.

Повод: /mnt/zajednicko/април.doc (стварна листа) је падала — заглавље у спојеним
ћелијама, ознака „Име и презиме:" носи placeholder а стварно име је у реду испод,
ред „Датум 1..31" није први ред. Фикстура `app_export_header_in_table` то
анонимизовано репродукује.
"""

import io
import os
import sys

import docx
import pytest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'tests', 'fixtures', 'radne_liste'))

import radna_lista_word_parser as rl
import make_fixtures as mk


# --- стварни распоред (заглавље унутар табеле) ---
def test_app_export_header_in_table_parsira():
    r = rl.parse_radna_lista(mk.build_app_export_header_in_table())
    assert r['employee_name'] == 'Петар Петровић'
    assert r['organization_unit'] == 'Геолошко одељење'
    assert r['position'] == 'Виши кустос'
    assert r['month'] == 9 and r['year'] == 2025
    assert r['orientation'] == 'days_cols'
    assert r['daily_data'] == mk.expected_daily_data()


def test_placeholder_u_oznaci_se_ignorise():
    """Ознака „Име и презиме: <placeholder>" не сме да прегази стварно име
    из реда испод."""
    r = rl.parse_radna_lista(mk.build_app_export_header_in_table())
    assert 'Ссс' not in r['employee_name']
    assert r['employee_name'] == 'Петар Петровић'


# --- конкретне поруке одбијања ---
def _doc_bytes(build):
    doc = docx.Document()
    build(doc)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def test_poruka_nema_tabele():
    def b(doc):
        doc.add_paragraph('Име и презиме: Петар Петровић')
        doc.add_paragraph('Септембар 2025')
    with pytest.raises(rl.RadnaListaParseError) as ei:
        rl.parse_radna_lista(_doc_bytes(b))
    assert 'нема ниједне табеле' in str(ei.value)


def test_poruka_dani_bez_kategorija():
    def b(doc):
        doc.add_paragraph('Име и презиме: Петар Петровић')
        doc.add_paragraph('Септембар 2025')
        t = doc.add_table(rows=1, cols=33)
        t.rows[0].cells[0].text = 'Датум'
        for d in range(1, 32):
            t.rows[0].cells[d].text = str(d)
        # редови без иједне препознате категорије
        for _ in range(3):
            t.add_row().cells[0].text = 'Непозната ставка'
    with pytest.raises(rl.RadnaListaParseError) as ei:
        rl.parse_radna_lista(_doc_bytes(b))
    assert 'дане' in str(ei.value) and 'категориј' in str(ei.value)


def test_poruka_kategorije_bez_dana():
    def b(doc):
        doc.add_paragraph('Име и презиме: Петар Петровић')
        doc.add_paragraph('Септембар 2025')
        t = doc.add_table(rows=1, cols=3)
        t.rows[0].cells[0].text = 'Категорија'
        # категорије у колони 0, али без реда са данима
        for lab in ('Рад на месту', 'Ван музеја', 'Годишњи одмор'):
            t.add_row().cells[0].text = lab
    with pytest.raises(rl.RadnaListaParseError) as ei:
        rl.parse_radna_lista(_doc_bytes(b))
    assert 'недостаје ред са данима' in str(ei.value)

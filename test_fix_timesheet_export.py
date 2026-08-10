#!/usr/bin/env python3
"""
Reproducing test for the timesheet Word export bug:
the page-2 "Обављени послови" section reads the wrong DB column
(extraordinary_tasks) instead of the active special_tasks column, so
the structured work description is always blank in the generated document.
"""

import os

os.environ.setdefault('FLASK_ENV', 'testing')
os.environ.setdefault('SECRET_KEY', 'test-secret')
os.environ.setdefault('REDIS_URL', '')
os.environ.setdefault('SESSION_TYPE', 'filesystem')
os.environ.setdefault('SESSION_FILE_DIR', '/tmp/museum-test-timesheet-export')

import json
import tempfile
import unittest
from contextlib import ExitStack, contextmanager
from unittest.mock import MagicMock, patch

import timesheet_word_export
from timesheet_word_export import DOCX_AVAILABLE, generate_word_document

if DOCX_AVAILABLE:
    from docx import Document as _RealDocument


WORK_JSON = json.dumps({
    'cat_2_1': 'Сређивање минералошке збирке',
    'cat_8': 'Остали послови у музеју',
})


def _make_header(with_special=True, with_extraordinary=False):
    header = {
        'id': 1,
        'employee_name': 'Петар Петровић',
        'month': 6,
        'year': 2025,
        'organization_unit': 'Минералошко одељење',
        'position': 'Кустос',
        'duties_summary': None,
        'created_at': None,
        'extraordinary_tasks': WORK_JSON if with_extraordinary else None,
    }
    if with_special:
        header['special_tasks'] = WORK_JSON
    return header


def _install_fake_psycopg(header, captured_sql):
    """Return a context manager that patches psycopg.connect with a fake
    connection. The header dict is returned by the first fetchone(); daily
    rows are empty. captured_sql collects executed SQL statements."""

    cur = MagicMock()

    def execute(sql, params=None):
        captured_sql.append(sql)

    cur.execute.side_effect = execute
    cur.fetchone.return_value = header
    cur.fetchall.return_value = []

    @contextmanager
    def cursor_cm():
        yield cur

    conn = MagicMock()
    conn.cursor.side_effect = cursor_cm

    @contextmanager
    def connect_cm(*args, **kwargs):
        yield conn

    return patch.object(timesheet_word_export.psycopg, 'connect', side_effect=connect_cm)


def _extract_text(path):
    doc = _RealDocument(path)
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return '\n'.join(texts)


@contextmanager
def _redirect_output(tmp_dir):
    """Redirect generated documents into a writable temp dir, since the
    repo's exports/timesheets dir is not writable in this environment."""

    from docx.document import Document as _DocClass
    real_save = _DocClass.save

    def patched_save(doc, path, *a, **k):
        target = os.path.join(tmp_dir, os.path.basename(path))
        return real_save(doc, target, *a, **k)

    # Document uses __slots__, so save is patched on the class, not the instance.
    with patch('timesheet_word_export.os.makedirs'), \
            patch.object(_DocClass, 'save', patched_save):
        yield


def _run_export(header):
    captured_sql = []
    with ExitStack() as stack:
        tmp_dir = stack.enter_context(tempfile.TemporaryDirectory())
        stack.enter_context(_install_fake_psycopg(header, captured_sql))
        stack.enter_context(_redirect_output(tmp_dir))
        generate_word_document(1, 'postgresql://x/y')
        produced = os.path.join(tmp_dir, os.listdir(tmp_dir)[0])
        text = _extract_text(produced)
    return text, captured_sql


@unittest.skipUnless(DOCX_AVAILABLE, 'python-docx not available')
class TestWordExportWorkDescription(unittest.TestCase):

    def test_special_tasks_rendered_on_page_two(self):
        """The structured work description stored in special_tasks must be
        rendered; the 'no data' placeholder must NOT appear."""
        header = _make_header(with_special=True, with_extraordinary=False)
        text, _ = _run_export(header)

        self.assertIn('Сређивање минералошке збирке', text)
        self.assertIn('Остали послови у музеју', text)
        self.assertNotIn('Нема података о обављеним пословима.', text)

    def test_select_fetches_special_tasks_column(self):
        """The header SELECT must request the special_tasks column."""
        header = _make_header(with_special=True, with_extraordinary=False)
        _, captured_sql = _run_export(header)

        joined = ' '.join(captured_sql)
        self.assertIn('special_tasks', joined)

    def test_extraordinary_tasks_legacy_fallback_still_works(self):
        """When only the legacy extraordinary_tasks column holds the JSON,
        it must still render (read precedence falls back to it)."""
        header = _make_header(with_special=False, with_extraordinary=True)
        text, _ = _run_export(header)

        self.assertIn('Сређивање минералошке збирке', text)
        self.assertNotIn('Нема података о обављеним пословима.', text)


if __name__ == '__main__':
    unittest.main()

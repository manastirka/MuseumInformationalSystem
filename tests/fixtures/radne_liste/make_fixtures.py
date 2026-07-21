#!/usr/bin/env python3
"""Генератор .docx фикстура за увоз радних листа из Word-а.

Прави мале .docx документе који опонашају стварну радну листу
(`Радна_листа_Др Александар Луковић_4_2015.docx`): једна табела-матрица
присуства + пасуси заглавља („Име и презиме", „Месец/Година",
„Организациона јединица", „Радно место") + одељак „Обављени послови".

Свака `build_*` функција враћа `io.BytesIO` спреман за
`radna_lista_word_parser.parse_radna_lista(...)` (које прима file-like).
Покренут као `__main__`, уписује .docx фајлове у овај фолдер за ручни преглед.

Изворни податак је ЈЕДАН речник (`SAMPLE_DAILY`) па и ћир/лат/транспоновани
документи носе ИСТУ матрицу — тестови пореде против `expected_daily_data()`.
"""

import io
import os

import docx


# ---------------------------------------------------------------------------
# Изворни подаци (дељени између свих оријентација/писама)
# ---------------------------------------------------------------------------
EMPLOYEE_NAME = 'Др Александар Луковић'
ORG_UNIT = 'Геолошко одељење'
POSITION = 'Виши кустос'
MONTH = 9
YEAR = 2025

# Дан → {канонски_код: сати}. Викенди (6,7,13,14,20,21,27,28) су празни.
SAMPLE_DAILY = {
    1: {'rad_na_mestu': 8.0},
    2: {'rad_na_mestu': 8.0},
    3: {'rad_na_mestu': 8.0},
    4: {'rad_na_mestu': 8.0},
    5: {'rad_na_mestu': 8.0},
    8: {'rad_na_mestu': 8.0},
    9: {'rad_na_mestu': 8.0},
    10: {'rad_na_mestu': 8.0},
    11: {'van_muzeja': 8.0},
    12: {'rad_na_mestu': 8.0},
    15: {'godisnji_odmor': 8.0},
    16: {'godisnji_odmor': 8.0},
    17: {'godisnji_odmor': 8.0},
    18: {'rad_na_mestu': 8.0},
    19: {'rad_na_mestu': 8.0},
    22: {'drzavni_praznik': 8.0},
    23: {'rad_na_mestu': 8.0},
    24: {'rad_na_mestu': 8.0},
    25: {'rad_na_mestu': 8.0},
    26: {'van_muzeja': 8.0},
    29: {'rad_na_mestu': 8.0},
    30: {'rad_na_mestu': 8.0},
}

# Мала, „сигурна" матрица (сви дани ≤ 28) — безбедна за било који месец,
# укључујући фебруар. Користи је ток ТЕКУЋЕ листе (default_entry_period).
SAFE_DAILY = {
    1: {'rad_na_mestu': 8.0},
    2: {'rad_na_mestu': 8.0},
    3: {'van_muzeja': 8.0},
    4: {'godisnji_odmor': 8.0},
    25: {'rad_na_mestu': 8.0},
}

# Канонски кодови у РЕДОСЛЕДУ колона матрице (исти као CATEGORY_CODES).
CATEGORY_CODES = [
    'rad_na_mestu', 'van_muzeja', 'godisnji_odmor', 'drzavni_praznik',
    'placeno_odsustvo', 'ostalo_odsustvo', 'bolovanje_manje_30', 'bolovanje_vece_30',
]

# Ознаке колона по писму (индекс се поклапа са CATEGORY_CODES).
CYR_LABELS = [
    'Рад на самом месту (у музеју)', 'Ван музеја', 'Годишњи одмор',
    'Државни празник', 'Плаћено одсуство', 'Остало одсуство',
    'Боловање < 30 дана', 'Боловање ≥ 30 дана',
]
LAT_LABELS = [
    'Rad na mestu', 'Van muzeja', 'Godišnji odmor', 'Državni praznik',
    'Plaćeno odsustvo', 'Ostalo odsustvo', 'Bolovanje < 30 dana',
    'Bolovanje ≥ 30 dana',
]

_WORK_DESCRIPTION = (
    'Редовни стручни послови кустоса минералога: документовање и '
    'каталогизација узорака, рад са збиркама, припрема изложби.'
)


# ---------------------------------------------------------------------------
# Помоћне
# ---------------------------------------------------------------------------
def expected_daily_data(daily=None):
    """Оно што парсер треба да врати: {str(dan): {code: hours}} сортирано."""
    src = SAMPLE_DAILY if daily is None else daily
    return {str(d): dict(rec) for d, rec in sorted(src.items())}


def _fmt_hours(h):
    return str(int(h)) if float(h) == int(h) else str(h)


def _to_bytes(doc):
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _add_header_paragraphs(doc, name, month, year, month_text=None,
                           labels_latin=False):
    doc.add_paragraph('Радна листа - Извештај' if not labels_latin
                      else 'Radna lista - Izveštaj')
    if labels_latin:
        if name:
            doc.add_paragraph(f'Ime i prezime: {name}')
        doc.add_paragraph(f'Mesec: {month_text or month}   Godina: {year}')
        doc.add_paragraph(f'Organizaciona jedinica: {ORG_UNIT}')
        doc.add_paragraph(f'Radno mesto: {POSITION}')
    else:
        if name:
            doc.add_paragraph(f'Име и презиме: {name}')
        doc.add_paragraph(f'Месец: {month_text or month}   Година: {year}')
        doc.add_paragraph(f'Организациона јединица: {ORG_UNIT}')
        doc.add_paragraph(f'Радно место: {POSITION}')


def _add_work_section(doc, labels_latin=False):
    doc.add_paragraph('Obavljeni poslovi' if labels_latin else 'Обављени послови')
    doc.add_paragraph(_WORK_DESCRIPTION)


def _add_matrix_days_rows(doc, labels, daily):
    """Дани као РЕДОВИ × категорије као КОЛОНЕ (изворни/апп формат)."""
    ncols = 1 + len(labels)
    table = doc.add_table(rows=1, cols=ncols)
    hdr = table.rows[0].cells
    hdr[0].text = 'Дан'
    for j, lab in enumerate(labels, start=1):
        hdr[j].text = lab
    for day in range(1, 32):
        rec = daily.get(day, {})
        cells = table.add_row().cells
        cells[0].text = str(day)
        for j, code in enumerate(CATEGORY_CODES, start=1):
            if code in rec:
                cells[j].text = _fmt_hours(rec[code])
    # УКУПНО ред (парсер га препознаје и прескаче).
    total = table.add_row().cells
    total[0].text = 'УКУПНО'
    return table


def _add_matrix_days_cols(doc, labels, daily):
    """Категорије као РЕДОВИ × дани као КОЛОНЕ (транспоновано)."""
    ncols = 1 + 31
    table = doc.add_table(rows=1, cols=ncols)
    hdr = table.rows[0].cells
    hdr[0].text = 'Категорија'
    for day in range(1, 32):
        hdr[day].text = str(day)
    for code, lab in zip(CATEGORY_CODES, labels):
        cells = table.add_row().cells
        cells[0].text = lab
        for day in range(1, 32):
            rec = daily.get(day, {})
            if code in rec:
                cells[day].text = _fmt_hours(rec[code])
    return table


# ---------------------------------------------------------------------------
# Јавни билдери
# ---------------------------------------------------------------------------
def build_correct_matrix_days_rows(name=EMPLOYEE_NAME, month=MONTH, year=YEAR,
                                   daily=None):
    """(a) Стварни распоред: ћирилица, дани редови × категорије колоне."""
    doc = docx.Document()
    _add_header_paragraphs(doc, name, month, year)
    _add_matrix_days_rows(doc, CYR_LABELS, SAMPLE_DAILY if daily is None else daily)
    _add_work_section(doc)
    return _to_bytes(doc)


def build_transposed_days_cols(name=EMPLOYEE_NAME, month=MONTH, year=YEAR,
                               daily=None):
    """(b) Иста матрица, али транспонована: категорије редови × дани колоне."""
    doc = docx.Document()
    _add_header_paragraphs(doc, name, month, year)
    _add_matrix_days_cols(doc, CYR_LABELS, SAMPLE_DAILY if daily is None else daily)
    _add_work_section(doc)
    return _to_bytes(doc)


_MONTH_NAMES_CYR = {
    1: 'јануар', 2: 'фебруар', 3: 'март', 4: 'април', 5: 'мај', 6: 'јун',
    7: 'јул', 8: 'август', 9: 'септембар', 10: 'октобар', 11: 'новембар',
    12: 'децембар',
}


def build_app_export_header_in_table(name='Петар Петровић', month=MONTH,
                                     year=YEAR, daily=None):
    """(f) СТВАРНИ апп-извоз (анонимизован по узору на /mnt/zajednicko/април.doc):
    заглавље је УНУТАР табеле (спојене ћелије), ознака „Име и презиме:" носи
    placeholder а стварно име је у реду ИСПОД; ред „Датум 1..31" није први ред
    (има преамбулу), категорије су редови (транспоновано). Овај распоред је
    раније рушио парсер."""
    daily = SAMPLE_DAILY if daily is None else daily
    doc = docx.Document()
    ncols = 1 + 31 + 1  # ознака + дани 1..31 + „Укупно:"
    table = doc.add_table(rows=0, cols=ncols)

    r0 = table.add_row().cells
    r0[0].text = 'Природњачки музеј у Београду'

    r1 = table.add_row().cells
    r1[0].text = 'Име и презиме: Сссссссссс'   # placeholder у ознаци (као у оригиналу)
    r1[9].text = f'РАДНА ЛИСТА за месец  {_MONTH_NAMES_CYR.get(month, month)} {year}.'
    r1[26].text = 'Организациона јединица:'

    r2 = table.add_row().cells
    r2[0].text = name                          # стварно име ИСПОД ознаке
    r2[9].text = f'Радно место: {POSITION}'
    r2[26].text = ORG_UNIT

    r3 = table.add_row().cells                 # ред-заглавље (није први ред!)
    r3[0].text = 'Датум'
    for day in range(1, 32):
        r3[day].text = str(day)
    r3[32].text = 'Укупно:'

    for code, lab in zip(CATEGORY_CODES, CYR_LABELS):
        cells = table.add_row().cells
        cells[0].text = lab
        for day in range(1, 32):
            if code in daily.get(day, {}):
                cells[day].text = _fmt_hours(daily[day][code])

    table.add_row().cells[0].text = 'Укупно:'

    doc.add_paragraph('Обављени послови')
    doc.add_paragraph(f'{_MONTH_NAMES_CYR.get(month, month).capitalize()} {year}.')
    doc.add_paragraph('1. Рад у збиркама: (анонимизован опис за тест).')
    return _to_bytes(doc)


def build_latin_variant(name='Petar Petrović', month=MONTH, year=YEAR,
                        daily=None):
    """(c) Латиничне ознаке + месец као име („Septembar 2025")."""
    doc = docx.Document()
    month_names = {
        1: 'Januar', 2: 'Februar', 3: 'Mart', 4: 'April', 5: 'Maj', 6: 'Jun',
        7: 'Jul', 8: 'Avgust', 9: 'Septembar', 10: 'Oktobar', 11: 'Novembar',
        12: 'Decembar',
    }
    month_text = f'{month_names[month]} {year}'
    _add_header_paragraphs(doc, name, month, year, month_text=month_text,
                           labels_latin=True)
    _add_matrix_days_rows(doc, LAT_LABELS, SAMPLE_DAILY if daily is None else daily)
    _add_work_section(doc, labels_latin=True)
    return _to_bytes(doc)


def build_broken_no_matrix(name=EMPLOYEE_NAME, month=MONTH, year=YEAR):
    """(d) Заглавље постоји, али НЕМА табеле-матрице."""
    doc = docx.Document()
    _add_header_paragraphs(doc, name, month, year)
    _add_work_section(doc)
    return _to_bytes(doc)


def build_broken_no_name(month=MONTH, year=YEAR, daily=None):
    """(e) Матрица постоји, месец/година постоје, али НЕМА имена запосленог."""
    doc = docx.Document()
    # Свесно изостављамо пасус „Име и презиме".
    doc.add_paragraph('Радна листа - Извештај')
    doc.add_paragraph(f'Месец: {month}   Година: {year}')
    doc.add_paragraph(f'Организациона јединица: {ORG_UNIT}')
    doc.add_paragraph(f'Радно место: {POSITION}')
    _add_matrix_days_rows(doc, CYR_LABELS, SAMPLE_DAILY if daily is None else daily)
    _add_work_section(doc)
    return _to_bytes(doc)


# Мапа име→билдер (за __main__ упис и лакше набрајање у тестовима).
CASES = {
    'correct_matrix_days_rows': build_correct_matrix_days_rows,
    'transposed_days_cols': build_transposed_days_cols,
    'latin_variant': build_latin_variant,
    'app_export_header_in_table': build_app_export_header_in_table,
    'broken_no_matrix': build_broken_no_matrix,
    'broken_no_name': build_broken_no_name,
}


def write_all(target_dir=None):
    """Упиши све фикстуре као .docx у дати фолдер (подразумевано овде)."""
    target_dir = target_dir or os.path.dirname(os.path.abspath(__file__))
    os.makedirs(target_dir, exist_ok=True)
    written = []
    for case_name, builder in CASES.items():
        buf = builder()
        path = os.path.join(target_dir, f'{case_name}.docx')
        with open(path, 'wb') as fh:
            fh.write(buf.getvalue())
        written.append(path)
    return written


if __name__ == '__main__':
    for p in write_all():
        print('написано:', p)

"""Парсер радне листе из Word (.docx) документа.

Flask-free и без базе — само python-docx + чиста логика, да га тестови и увозни
токови зову директно. Толерантан на варијације:
  * оријентација матрице присуства: дани као РЕДОВИ × категорије као КОЛОНЕ
    (изворни/апп формат) ИЛИ категорије као редови × дани као колоне;
  * ћирилица/латиница у ознакама категорија и месецима;
  * заглавље у пасусима ИЛИ у табели.

Никад тихо не погађа: ако не разуме документ (нема запосленог, месеца/године,
или матрице) диже `RadnaListaParseError` са јасним разлогом; мање недоумице
скупља у `warnings`.
"""

import re
import unicodedata
from datetime import date


class RadnaListaParseError(Exception):
    """Документ се не може поуздано рашчланити (порука објашњава зашто)."""


# Канонски кодови категорија (исти као timesheet_report_days / frontend).
CATEGORY_CODES = [
    'rad_na_mestu', 'van_muzeja', 'godisnji_odmor', 'drzavni_praznik',
    'placeno_odsustvo', 'ostalo_odsustvo', 'bolovanje_manje_30', 'bolovanje_vece_30',
]
# Категорије где ненумерички текст у ћелији (нпр. име празника) значи цео дан.
_FULL_DAY_CATEGORIES = {
    'godisnji_odmor', 'drzavni_praznik', 'placeno_odsustvo',
    'ostalo_odsustvo', 'bolovanje_manje_30', 'bolovanje_vece_30',
}
_STANDARD_DAY_HOURS = 8.0

# Месеци: ћир + лат, номинатив/локатив основе → број.
_MONTHS = {
    1: ('јануар', 'januar'), 2: ('фебруар', 'februar'), 3: ('март', 'mart'),
    4: ('април', 'april'), 5: ('мај', 'maj'), 6: ('јун', 'jun'),
    7: ('јул', 'jul'), 8: ('август', 'avgust', 'аугуст'), 9: ('септембар', 'septembar'),
    10: ('октобар', 'oktobar'), 11: ('новембар', 'novembar'), 12: ('децембар', 'decembar'),
}


def _norm(s):
    """Мала слова + сажети размаци + фолдовање латиничних дијакритика
    (č→c, ž→z, š→s, ć→c, đ→d) преко NFKD. Ћирилица остаје нетакнута (српска
    ћирилична слова немају канонску декомпозицију), па „плаћено" (ћир) и
    „plaćeno" (лат) оба падну на исте кључне речи."""
    s = unicodedata.normalize('NFKD', str(s or ''))
    s = ''.join(c for c in s if not unicodedata.combining(c))
    return re.sub(r'\s+', ' ', s.strip().lower())


def _match_category(label):
    """Ознаку колоне/реда мапирај на канонски код категорије или None."""
    t = _norm(label)
    if not t:
        return None
    has = lambda *ks: any(k in t for k in ks)
    if has('боловањ', 'bolovanj'):
        if has('>30', '> 30', 'веће', 'vece', 'више', 'vise', 'преко', 'preko', 'дуже', 'duze'):
            return 'bolovanje_vece_30'
        if has('<30', '< 30', 'мање', 'manje', 'до 30', 'do 30', 'краће', 'krace'):
            return 'bolovanje_manje_30'
        return 'bolovanje_manje_30'  # неозначено боловање → до 30
    if has('годишњ', 'godisnj') or (has('одмор', 'odmor') and not has('плаћен')):
        return 'godisnji_odmor'
    if has('државн', 'drzavn', 'празник', 'praznik'):
        return 'drzavni_praznik'
    if has('плаћен', 'placen'):
        return 'placeno_odsustvo'
    if has('остал', 'ostal'):
        return 'ostalo_odsustvo'
    if has('ван муз', 'van muz') or (has('ван', 'van') and has('муз', 'muz')):
        return 'van_muzeja'
    if has('рад') or has('rad') or has('у музеј', 'u muzej', 'на месту', 'na mestu'):
        return 'rad_na_mestu'
    return None


def _as_int_day(text):
    t = _norm(text)
    m = re.fullmatch(r'(\d{1,2})\.?', t)
    if m:
        d = int(m.group(1))
        return d if 1 <= d <= 31 else None
    return None


def _is_total_label(text):
    return any(k in _norm(text) for k in ('укупно', 'ukupno', 'збир', 'zbir', 'σ'))


def _cell_hours(text, category):
    """Сати из ћелије: број → сати; X/Х → 0 (викенд/нерадно); ненумерички текст
    у категорији одсуства → цео дан (8); иначе 0."""
    t = (text or '').strip()
    if not t or t.upper() in ('X', 'Х', '/', '-', '—'):
        return 0.0
    m = re.search(r'(\d+(?:[.,]\d+)?)', t)
    if m:
        try:
            return float(m.group(1).replace(',', '.'))
        except ValueError:
            return 0.0
    if category in _FULL_DAY_CATEGORIES:
        return _STANDARD_DAY_HOURS
    return 0.0


# ---------------------------------------------------------------------------
# Заглавље
# ---------------------------------------------------------------------------
def _collect_label_values(doc):
    """Скупи „ознака: вредност" парове из пасуса И ћелија табела."""
    pairs = []
    texts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                texts.append(c.text)
    for raw in texts:
        for line in str(raw).split('\n'):
            if ':' in line:
                k, _, v = line.partition(':')
                pairs.append((_norm(k), v.strip()))
    return pairs, texts


_NAME_KEYS = ('име и презиме', 'ime i prezime', 'запослени', 'zaposleni')
_ORG_KEYS = ('организацион', 'organizacion', 'одељењ', 'odeljen', 'јединиц', 'jedinic')
_POS_KEYS = ('радно место', 'radno mesto')


def _extract_header(doc, matrix_table=None, header_row=0):
    pairs, texts = _collect_label_values(doc)

    def find(*keys):
        for k, v in pairs:
            if any(key in k for key in keys) and v:
                return v.strip()
        return None

    name = find(*_NAME_KEYS)
    org = find(*_ORG_KEYS)
    position = find(*_POS_KEYS)

    # Заглавље може бити УНУТАР матрице (спојене ћелије): нпр. апп-извоз држи
    # ознаку „Име и презиме:" у једном реду, а стварну вредност у реду ИСПОД
    # (ознака чак носи placeholder). Тада вадимо структурно (ћелија испод ознаке).
    if matrix_table is not None and header_row > 0:
        name = _value_by_label(matrix_table, _NAME_KEYS, prefer_below=True) or name
        org = _value_by_label(matrix_table, _ORG_KEYS, prefer_below=True) or org
        position = _value_by_label(matrix_table, _POS_KEYS, prefer_below=False) or position

    month, year = _extract_month_year(pairs, texts)
    return {
        'employee_name': _clean_name(name),
        'organization_unit': org or None,
        'position': position or None,
        'month': month,
        'year': year,
    }


def _value_by_label(table, keys, prefer_below):
    """Вредност уз ОЗНАКУ у табели. Ознака је ћелија где се кључ појављује ПРЕ
    двотачке („Организациона јединица: …") — тиме се вредносне ћелије које
    случајно садрже кључну реч (нпр. „Геолошко одељење") НЕ третирају као ознаке.
    prefer_below=True → узми ћелију ИСПОД ознаке (вредност је у реду испод, ознака
    може носити placeholder); иначе део после двотачке у истој ћелији."""
    rows = table.rows
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row.cells):
            raw = cell.text
            if ':' not in raw:
                continue
            label_part = _norm(raw.split(':', 1)[0])
            if not any(k in label_part for k in keys):
                continue
            inline = raw.split(':', 1)[1].strip()
            below = ''
            if ri + 1 < len(rows) and ci < len(rows[ri + 1].cells):
                below = rows[ri + 1].cells[ci].text.strip()
            if below.endswith(':'):   # ред испод је опет ознака, не вредност
                below = ''
            if prefer_below:
                if below:
                    return below
                if inline and not _looks_placeholder(inline):
                    return inline
            else:
                if inline:
                    return inline
                if below:
                    return below
    return None


def _looks_placeholder(text):
    """Груба детекција placeholder-а (нпр. „СссссSSSSSsssssa"): нема размака и
    има мешавину случајних понављања/оба писма без смисла."""
    t = text.strip()
    if not t or ' ' in t:
        return False
    # понављање истог слова 3+ пута заредом → placeholder
    return bool(re.search(r'(.)\1\1', t))


def _clean_name(name):
    if not name:
        return None
    # скини титуле/суфиксе типа „Др ", вишеструке размаке
    n = re.sub(r'\s+', ' ', name.strip())
    return n or None


def _extract_month_year(pairs, texts):
    month = year = None
    # 1) експлицитно „Месец: 4  Година: 2015"
    for k, v in pairs:
        if 'месец' in k or 'mesec' in k or 'период' in k or 'period' in k:
            m, y = _parse_month_year_text(v)
            month = month or m
            year = year or y
        if ('година' in k or 'godina' in k) and year is None:
            ym = re.search(r'(19|20)\d{2}', v)
            if ym:
                year = int(ym.group(0))
    # 2) слободан текст „РАДНА ЛИСТА за месец август 2025" / „Август 2025"
    if month is None or year is None:
        for raw in texts:
            m, y = _parse_month_year_text(raw)
            if month is None and m:
                month = m
            if year is None and y:
                year = y
            if month and year:
                break
    return month, year


def _parse_month_year_text(text):
    t = _norm(text)
    month = None
    for num, names in _MONTHS.items():
        if any(nm in t for nm in names):
            month = num
            break
    if month is None:
        m = re.search(r'месец\D{0,4}(\d{1,2})|mesec\D{0,4}(\d{1,2})', t)
        if m:
            val = int(m.group(1) or m.group(2))
            if 1 <= val <= 12:
                month = val
    ym = re.search(r'(19|20)\d{2}', t)
    year = int(ym.group(0)) if ym else None
    return month, year


# ---------------------------------------------------------------------------
# Матрица присуства
# ---------------------------------------------------------------------------
def _table_grid(table):
    return [[c.text for c in row.cells] for row in table.rows]


def _find_matrix(doc):
    """Нађи табелу матрице. Толерантан на ПРЕАМБУЛУ (заглавље унутар табеле,
    спојене ћелије): ред-заглавље (дани ИЛИ категорије) се тражи БИЛО ГДЕ, не
    само у првом реду. Враћа (table, grid, orientation, header_row) или None.

    orientation 'days_rows' = дани редови × категорије колоне; 'days_cols' =
    категорије редови × дани колоне. header_row = индекс реда-заглавља.
    """
    best = None
    for t in doc.tables:
        grid = _table_grid(t)
        if len(grid) < 2 or len(grid[0]) < 2:
            continue
        for hr in range(len(grid)):
            # days_cols: ред са данима (Датум 1..31), категорије у колони 0 испод
            days_in_row = sum(1 for c in grid[hr][1:] if _as_int_day(c))
            if days_in_row >= 8:
                cats_below = sum(1 for r in range(hr + 1, len(grid))
                                 if _match_category(grid[r][0]))
                if cats_below >= 2:
                    best = _better(best, (days_in_row + cats_below, t, grid, 'days_cols', hr))
            # days_rows: ред са ≥2 категорије, дани у колони 0 испод
            cats_in_row = sum(1 for c in grid[hr] if _match_category(c))
            if cats_in_row >= 2:
                days_below = sum(1 for r in range(hr + 1, len(grid))
                                 if _as_int_day(grid[r][0]))
                if days_below >= 8:
                    best = _better(best, (cats_in_row + days_below, t, grid, 'days_rows', hr))
    if best is None:
        return None
    return best[1], best[2], best[3], best[4]


def _better(best, cand):
    return cand if (best is None or cand[0] > best[0]) else best


def _matrix_error_reason(doc):
    """Конкретан разлог зашто матрица није призната — да корисник/админ зна шта
    фали без дебуговања."""
    if not doc.tables:
        return ('нема ниједне табеле у документу — недостаје матрица присуства '
                '(табела са данима 1–31 и редовима категорија).')
    for t in doc.tables:
        grid = _table_grid(t)
        has_days = any(sum(1 for c in row[1:] if _as_int_day(c)) >= 8 for row in grid) \
            or sum(1 for r in range(len(grid)) if _as_int_day(grid[r][0])) >= 8
        has_cats = any(_match_category(c) for row in grid for c in row)
        if has_days and not has_cats:
            return ('табела има дане, али ниједан ред није препознат као категорија '
                    '(Рад на месту, Ван музеја, Годишњи одмор, Државни празник, '
                    'Плаћено/Остало одсуство, Боловање) — можда другачији називи.')
        if has_cats and not has_days:
            return ('табела има категорије, али недостаје ред са данима (1–31) / „Датум".')
    return ('није пронађена матрица присуства (табела са данима 1–31 и '
            'редовима категорија).')


def _parse_matrix(grid, orientation, header_row):
    """Врати daily_data {day: {cat: hours}} само за дане са неким сатима.
    Редови ИСПОД header_row су подаци (изнад је преамбула/заглавље)."""
    daily = {}
    header = grid[header_row]
    if orientation == 'days_rows':
        col_cat = {j: _match_category(header[j]) for j in range(1, len(header))}
        for row in grid[header_row + 1:]:
            if not row or _is_total_label(row[0]):
                continue
            day = _as_int_day(row[0])
            if day is None:
                continue
            rec = {}
            for j, cat in col_cat.items():
                if cat and j < len(row):
                    h = _cell_hours(row[j], cat)
                    if h:
                        rec[cat] = rec.get(cat, 0.0) + h
            if rec:
                daily[day] = rec
    else:  # days_cols: категорије редови × дани колоне
        col_day = {j: _as_int_day(header[j]) for j in range(1, len(header))}
        for row in grid[header_row + 1:]:
            if not row:
                continue
            cat = _match_category(row[0])
            if not cat or _is_total_label(row[0]):
                continue
            for j, day in col_day.items():
                if day is None or j >= len(row):
                    continue
                h = _cell_hours(row[j], cat)
                if h:
                    daily.setdefault(day, {})
                    daily[day][cat] = daily[day].get(cat, 0.0) + h
    return daily


# ---------------------------------------------------------------------------
# Обављени послови
# ---------------------------------------------------------------------------
def _extract_work_description(doc):
    paras = [p.text for p in doc.paragraphs]
    out, grab = [], False
    for p in paras:
        norm = _norm(p)
        if not grab and ('обављени послови' in norm or 'obavljeni poslovi' in norm
                         or 'опис обављених' in norm or 'редовни послови' in norm):
            grab = True
            # прескочи сам наслов ако је чист наслов
            if norm in ('обављени послови', 'obavljeni poslovi', 'опис обављених послова'):
                continue
        if grab:
            if _norm(p) in ('потписи', 'potpisi', 'дневни уноси', 'dnevni unosi'):
                break
            if p.strip():
                out.append(p.strip())
    text = '\n'.join(out).strip()
    return text or None


# ---------------------------------------------------------------------------
# Јавни улаз
# ---------------------------------------------------------------------------
def _read_source_bytes(source):
    if hasattr(source, 'read'):
        try:
            source.seek(0)
        except Exception:
            pass
        return source.read()
    with open(source, 'rb') as fh:
        return fh.read()


def parse_radna_lista(source):
    """Рашчлани радну листу из .docx ИЛИ .doc (путања или file-like).

    Формат се одлучује по magic бајтовима (не екстензији): .docx се чита директно,
    стари .doc се конвертује кроз LibreOffice. Враћа dict или диже
    RadnaListaParseError са јасним разлогом.
    """
    import io
    import doc_konverzija
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise RadnaListaParseError(f'python-docx није доступан: {exc}')

    data = _read_source_bytes(source)
    try:
        docx_bytes = doc_konverzija.ensure_docx_bytes(data)
    except doc_konverzija.SofficeUnavailable:
        raise RadnaListaParseError(
            'Увоз .doc формата тренутно није могућ — обратите се администратору.')
    except doc_konverzija.DocConversionError as exc:
        raise RadnaListaParseError(f'.doc се није могао конвертовати у .docx: {exc}')
    except ValueError as desc:
        raise RadnaListaParseError(f'датотека није ни .doc ни .docx ({desc}).')

    try:
        doc = docx.Document(io.BytesIO(docx_bytes))
    except Exception as exc:
        raise RadnaListaParseError(f'датотека није исправан .docx: {exc}')

    warnings = []
    # Матрицу тражимо ПРЕ заглавља: тако знамо да ли је заглавље унутар табеле
    # (спојене ћелије) па вредности вадимо структурно.
    found = _find_matrix(doc)
    if found is None:
        raise RadnaListaParseError(_matrix_error_reason(doc))
    matrix_table, grid, orientation, header_row = found

    header = _extract_header(doc, matrix_table, header_row)
    if not header['employee_name']:
        raise RadnaListaParseError(
            'није пронађено име запосленог (заглавље „Име и презиме" није нађено '
            'ни у пасусима ни у табели).')
    if not header['month'] or not header['year']:
        raise RadnaListaParseError(
            'није препознат месец и/или година (нема „РАДНА ЛИСТА за месец … '
            '<година>" ни назива месеца).')
    if not 1 <= header['month'] <= 12:
        raise RadnaListaParseError(f'нелогичан месец: {header["month"]}.')

    daily = _parse_matrix(grid, orientation, header_row)
    if not daily:
        warnings.append('матрица је нађена, али ниједан дан нема унете сате.')

    totals = _totals(daily)
    if totals['radni'] == 0 and totals['sve'] == 0:
        warnings.append('укупан број сати је 0 — проверите документ.')

    return {
        'employee_name': header['employee_name'],
        'organization_unit': header['organization_unit'],
        'position': header['position'],
        'month': header['month'],
        'year': header['year'],
        'daily_data': {str(d): rec for d, rec in sorted(daily.items())},
        'work_description': _extract_work_description(doc),
        'totals': totals,
        'orientation': orientation,
        'warnings': warnings,
    }


def _totals(daily):
    radni = sum(rec.get('rad_na_mestu', 0) + rec.get('van_muzeja', 0)
                for rec in daily.values())
    sve = sum(sum(rec.values()) for rec in daily.values())
    per_cat = {c: 0.0 for c in CATEGORY_CODES}
    for rec in daily.values():
        for c, h in rec.items():
            per_cat[c] = per_cat.get(c, 0.0) + h
    return {'radni': radni, 'sve': sve, 'po_kategoriji': per_cat}
